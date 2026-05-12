"""
app/api/v1/endpoints/in_bang_ke.py
===================================
API Endpoints in bảng kê cá nhân (phiếu đánh giá + bảng kê công việc).

Endpoints:
1. GET /api/v1/in-bang-ke/phieu-danh-gia/{thang}/{nam} - In phiếu đánh giá tháng (PL-01A hoặc PL-01B)
2. GET /api/v1/in-bang-ke/bang-ke-cong-viec/{thang}/{nam} - In bảng kê công việc tháng (PL-02)
3. GET /api/v1/in-bang-ke/phieu-danh-gia-quy/{quy}/{nam} - In phiếu đánh giá quý (PL-01A hoặc PL-01B)
4. GET /api/v1/in-bang-ke/bang-ke-cong-viec-quy/{quy}/{nam} - In bảng kê công việc quý (PL-02)

Output: DOCX file download

Version: 1.1.0 (16/04/2026)
"""

import io
import logging
from pathlib import Path
from typing import Optional
from datetime import date, datetime
from decimal import Decimal
from uuid import UUID

from fastapi import APIRouter, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select, and_, func
from sqlalchemy.orm import selectinload

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from app.api.deps import DatabaseDep, ActiveUserDep
from app.models.user_org import CapBacVaiTro, CongChuc, VaiTro
from app.models.kpi_submission import KeKhaiCongViec, TrangThaiKeKhai
from app.models.kpi_assessment import (
    DanhGiaThang,
    LanhDaoChiSo,
    TieuChiChungDanhGia,
    TrangThaiTieuChi,
)
from app.models.leader_kpi import KeKhaiLanhDao, TrangThaiKeKhaiLD, TrangThaiHoanThanh
from app.models.phieu_danh_gia import (
    PhieuDanhGiaQuy,
    PhieuDanhGiaThang,
    TrangThaiPhieuDanhGia,
)
from app.models.task_catalog import DanhMucSpCongViec

logger = logging.getLogger(__name__)

router = APIRouter()

TEMPLATES_DIR = Path(__file__).parent.parent.parent.parent.parent / "templates"


# =============================================================================
# HELPER: SET FONT TIMES NEW ROMAN
# =============================================================================

def set_times_new_roman(run):
    """
    Set font Times New Roman cho run (bao gồm tiếng Việt có dấu).
    Phải set cả eastAsia để áp dụng cho tiếng Việt.
    """
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')


def set_cell_font_times_new_roman(cell):
    """
    Set font Times New Roman cho tất cả runs trong 1 cell.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_times_new_roman(run)


def replace_in_runs(paragraph, old_text: str, new_text: str) -> bool:
    """
    Replace text trong runs của paragraph, giữ nguyên cấu trúc runs khác.
    Tránh bug: khi template có nhiều runs (equation, subscript/superscript),
    việc replace toàn bộ paragraph.text sẽ làm mất format.

    Strategy:
    - old_text có thể bị tách ra nhiều runs (VD: "……….(%)" = run[6]="………." + run[7-9]="(%)")
    - Trước tiên tìm run chứa phần chính (dấu chấm placeholder)
    - Replace phần đó thành giá trị mới
    - Nếu old_text có phần tail (VD: "(%)", ".)"), tìm và clear các run chứa tail

    Returns:
        True nếu đã replace, False nếu không tìm thấy old_text
    """
    replaced = False

    # Case 1: old_text hoàn toàn trong 1 run (easy case)
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            set_times_new_roman(run)
            return True

    # Case 2: old_text bị tách ra nhiều runs (complex case)
    # VD: "……….(%)" → run[6]="): ………." + run[7-9]="(%)"
    # Strategy: tìm run chứa dấu chấm (……….hoặc …………), replace thành new_text, clear run chứa "(%)hoặc ."

    # Xác định placeholder core (dấu chấm) và tail ((...) hoặc.)
    # Các placeholder trong template:
    # - "……….(%)" → core="……….", tail="(%)"
    # - "…………………………." → core="…………………………..", tail=""

    # Tách old_text thành core và tail
    if "(%)" in old_text:
        core = old_text.replace("(%)", "")
        tail = "(%)"
    elif old_text.endswith("."):
        core = old_text
        tail = ""
    else:
        core = old_text
        tail = ""

    # Tìm run chứa core hoặc một phần của core
    for i, run in enumerate(paragraph.runs):
        if core in run.text:
            # Replace core thành new_text
            run.text = run.text.replace(core, new_text)
            set_times_new_roman(run)
            replaced = True

            # Nếu có tail, tìm và clear các run chứa tail
            if tail:
                for j in range(i+1, len(paragraph.runs)):
                    next_run = paragraph.runs[j]
                    if tail in next_run.text:
                        next_run.text = next_run.text.replace(tail, "")
                        break
                    # Nếu tail bị tách thành nhiều run (VD: "(", "%", ")"), clear từng phần
                    elif any(char in next_run.text for char in tail):
                        next_run.text = ""

            return True

        # Nếu run chứa dấu chấm (placeholder bị tách ra nhiều runs)
        # VD: "…………………………….." → run[2]="…………" + run[3]="……………….."
        # Strategy: tìm run đầu tiên chứa dấu "…" → replace thành new_text, clear các run tiếp theo chứa "…"
        elif "…" in run.text:
            # Replace dấu chấm thành new_text
            run.text = run.text.replace("…", "").replace(".", "") + new_text
            set_times_new_roman(run)
            replaced = True

            # Clear các run tiếp theo chứa dấu "…" hoặc "."
            for j in range(i+1, len(paragraph.runs)):
                next_run = paragraph.runs[j]
                if "…" in next_run.text or (next_run.text and next_run.text.strip() == "."):
                    next_run.text = ""

            # Nếu có tail (VD: "(%)", clear luôn)
            if tail:
                for j in range(i+1, len(paragraph.runs)):
                    next_run = paragraph.runs[j]
                    if any(char in next_run.text for char in tail):
                        next_run.text = ""

            return True

    return replaced


# =============================================================================
# HELPER: REPLACE PLACEHOLDER IN DOCX
# =============================================================================

def replace_placeholder_in_docx(doc: Document, placeholder: str, value: str):
    """
    Replace {{placeholder}} trong tất cả paragraphs của document.
    Áp dụng font Times New Roman cho text mới.
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Merge all runs để tránh bị split text
            full_text = para.text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full_text.replace(placeholder, value)
                set_times_new_roman(para.runs[0])
            else:
                para.text = full_text.replace(placeholder, value)

    # Replace trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        full_text = para.text
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = full_text.replace(placeholder, value)
                            set_times_new_roman(para.runs[0])
                        else:
                            para.text = full_text.replace(placeholder, value)


# =============================================================================
# HELPERS cho phiếu tháng & quý (v4.1.0)
# =============================================================================

DIA_DIEM_MAC_DINH = "Quảng Ninh"


def _format_ngay_dia_diem(dt: Optional[datetime]) -> str:
    """'Quảng Ninh, ngày DD tháng MM năm YYYY' — dùng dt nếu có, mặc định hôm nay."""
    d = dt.date() if dt else date.today()
    return f"{DIA_DIEM_MAC_DINH}, ngày {d.day:02d} tháng {d.month:02d} năm {d.year}"


def _fill_footer_date(doc: Document, dia_diem_str: str) -> None:
    """Thay '....., ngày.... tháng.... năm.....' trong footer (Table cuối)."""
    if len(doc.tables) < 3:
        return
    tbl = doc.tables[2]
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                t = p.text
                if "ngày" in t and "tháng" in t and "năm" in t and "....." in t:
                    for run in p.runs:
                        run.text = ""
                    if p.runs:
                        p.runs[0].text = dia_diem_str
                        set_times_new_roman(p.runs[0])
                    else:
                        r = p.add_run(dia_diem_str)
                        set_times_new_roman(r)


def _them_ten_ky(cell, ten: str, bold: bool = True, italic: bool = False) -> None:
    """
    Ghi tên ký ngay sau paragraph '(Ký, ghi rõ họ tên)' — ưu tiên dùng blank
    paragraph kế tiếp (nếu có) để tránh bị đẩy xuống cuối cell.

    Lý do: Cell trái (CC) thường có 8 blank paragraph dùng làm khoảng trắng
    chữ ký. Nếu append vào cuối thì tên rơi sâu ~8 dòng so với cell phải
    → lệch nhau. Ghi vào blank đầu tiên sau '(Ký…)' thì tên nằm sát header.
    """
    # Tìm index paragraph chứa '(Ký'
    idx_ky = -1
    for i, p in enumerate(cell.paragraphs):
        if "Ký" in p.text and "ghi rõ" in p.text:
            idx_ky = i
            break

    target_idx = idx_ky + 1 if idx_ky >= 0 else len(cell.paragraphs)

    if 0 <= target_idx < len(cell.paragraphs):
        target = cell.paragraphs[target_idx]
        # Clear old runs
        for run in target.runs:
            run.text = ""
        target.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if target.runs:
            target.runs[0].text = ten
            target.runs[0].bold = bold
            target.runs[0].italic = italic
            set_times_new_roman(target.runs[0])
        else:
            run = target.add_run(ten)
            run.bold = bold
            run.italic = italic
            set_times_new_roman(run)
    else:
        # Cell không có blank paragraph phía sau → append mới
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ten)
        run.bold = bold
        run.italic = italic
        set_times_new_roman(run)


def _fill_footer_signers(
    doc: Document,
    ten_cong_chuc: str,
    ten_nguoi_duyet: Optional[str],
) -> None:
    """Chèn tên CC (trái) và tên người duyệt (phải) phía dưới '(Ký, ghi rõ họ tên)'."""
    if len(doc.tables) < 3:
        return
    row = doc.tables[2].rows[0]
    if len(row.cells) >= 1 and ten_cong_chuc:
        _them_ten_ky(row.cells[0], ten_cong_chuc)
    if len(row.cells) >= 2 and ten_nguoi_duyet:
        _them_ten_ky(row.cells[1], ten_nguoi_duyet)


def _fill_section_after_heading(doc: Document, heading_prefix: str, content: str) -> bool:
    """Tìm paragraph bắt đầu bằng heading_prefix, ghi content vào paragraph trống kế tiếp."""
    if not content:
        return False
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip().startswith(heading_prefix):
            if i + 1 < len(paras):
                target = paras[i + 1]
                for run in target.runs:
                    run.text = ""
                if target.runs:
                    target.runs[0].text = content
                    set_times_new_roman(target.runs[0])
                else:
                    r = target.add_run(content)
                    set_times_new_roman(r)
                return True
    return False


async def _tim_nguoi_duyet_auto(db, cc: CongChuc) -> Optional[CongChuc]:
    """
    Xác định người duyệt mặc định theo cấp bậc CC (dùng cho auto-fill tên
    'Cấp có thẩm quyền' khi phiếu chưa được duyệt hoặc phiếu tháng).

    - CONG_CHUC / PHO_DON_VI             → TRUONG_DON_VI cùng đơn vị
    - TRUONG_DON_VI / PHO_CHI_CUC_TRUONG → CHI_CUC_TRUONG
    - CHI_CUC_TRUONG                     → None (tự ký)
    """
    cb = cc.vai_tro.cap_bac if cc.vai_tro else None
    if cb in (CapBacVaiTro.CONG_CHUC, CapBacVaiTro.PHO_DON_VI):
        stmt = (
            select(CongChuc)
            .join(VaiTro, VaiTro.id == CongChuc.vai_tro_id)
            .where(
                CongChuc.don_vi_id == cc.don_vi_id,
                CongChuc.is_active == True,
                VaiTro.cap_bac == CapBacVaiTro.TRUONG_DON_VI,
            )
            .limit(1)
        )
        return (await db.execute(stmt)).scalar_one_or_none()
    if cb in (CapBacVaiTro.TRUONG_DON_VI, CapBacVaiTro.PHO_CHI_CUC_TRUONG):
        stmt = (
            select(CongChuc)
            .join(VaiTro, VaiTro.id == CongChuc.vai_tro_id)
            .where(
                CongChuc.is_active == True,
                VaiTro.cap_bac == CapBacVaiTro.CHI_CUC_TRUONG,
            )
            .limit(1)
        )
        return (await db.execute(stmt)).scalar_one_or_none()
    return None


async def _lay_phieu_quy(
    db, cong_chuc_id, quy: int, nam: int
) -> Optional[PhieuDanhGiaQuy]:
    """Lấy phiếu đánh giá quý của CC (None nếu chưa tạo)."""
    stmt = (
        select(PhieuDanhGiaQuy)
        .options(selectinload(PhieuDanhGiaQuy.nguoi_phe_duyet))
        .where(PhieuDanhGiaQuy.cong_chuc_id == cong_chuc_id)
        .where(PhieuDanhGiaQuy.quy == quy)
        .where(PhieuDanhGiaQuy.nam == nam)
    )
    return (await db.execute(stmt)).scalar_one_or_none()


async def _lay_phieu_thang(
    db, cong_chuc_id, thang: int, nam: int
) -> Optional[PhieuDanhGiaThang]:
    """Lấy phiếu đánh giá tháng của CC (None nếu chưa tạo)."""
    stmt = (
        select(PhieuDanhGiaThang)
        .options(selectinload(PhieuDanhGiaThang.nguoi_phe_duyet))
        .where(PhieuDanhGiaThang.cong_chuc_id == cong_chuc_id)
        .where(PhieuDanhGiaThang.thang == thang)
        .where(PhieuDanhGiaThang.nam == nam)
    )
    return (await db.execute(stmt)).scalar_one_or_none()


async def _apply_auto_fill_chung(
    doc: Document,
    db,
    cong_chuc: CongChuc,
    ngay_ky: Optional[datetime] = None,
    ten_nguoi_duyet_override: Optional[str] = None,
) -> None:
    """
    Áp dụng auto-fill chung cho cả phiếu tháng & quý:
    - Footer: địa điểm + ngày
    - Bottom-left: tên CC
    - Bottom-right: tên TDV/CCT (auto hoặc override)
    """
    _fill_footer_date(doc, _format_ngay_dia_diem(ngay_ky))

    if ten_nguoi_duyet_override is not None:
        ten_nguoi_duyet = ten_nguoi_duyet_override
    else:
        nd = await _tim_nguoi_duyet_auto(db, cong_chuc)
        ten_nguoi_duyet = nd.ho_ten if nd else ""

    _fill_footer_signers(doc, cong_chuc.ho_ten or "", ten_nguoi_duyet)


# =============================================================================
# ENDPOINT 1: IN PHIẾU ĐÁNH GIÁ (PL-01A hoặc PL-01B)
# =============================================================================

@router.get("/phieu-danh-gia/{thang}/{nam}")
async def export_phieu_danh_gia(
    thang: int,
    nam: int,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """
    Xuất phiếu đánh giá cá nhân (PL-01A hoặc PL-01B) cho tháng/năm chỉ định.

    Auto-detect is_lanh_dao của user hiện tại:
    - is_lanh_dao = False → dùng template PL-01A (CC không lãnh đạo)
    - is_lanh_dao = True → dùng template PL-01B (Lãnh đạo)

    Args:
        thang: Tháng (1-12)
        nam: Năm (>= 2020)
        db: Database session
        current_user: User hiện tại (từ JWT)

    Returns:
        StreamingResponse: File DOCX download
    """
    # Validation
    if not (1 <= thang <= 12):
        raise HTTPException(status_code=400, detail="Tháng không hợp lệ (1-12)")
    if nam < 2020 or nam > 2100:
        raise HTTPException(status_code=400, detail="Năm không hợp lệ")

    # Load user với relationships
    stmt = (
        select(CongChuc)
        .options(
            selectinload(CongChuc.don_vi),
            selectinload(CongChuc.vai_tro),
        )
        .where(CongChuc.id == current_user.id)
    )
    result = await db.execute(stmt)
    cc = result.scalar_one_or_none()
    if not cc:
        raise HTTPException(status_code=404, detail="Không tìm thấy thông tin công chức")

    # Lấy dữ liệu đánh giá tháng
    stmt_dg = (
        select(DanhGiaThang)
        .where(
            and_(
                DanhGiaThang.cong_chuc_id == cc.id,
                DanhGiaThang.thang == thang,
                DanhGiaThang.nam == nam,
            )
        )
    )
    result_dg = await db.execute(stmt_dg)
    danh_gia = result_dg.scalar_one_or_none()

    # Lấy tiêu chí chung (phải JOIN qua DanhGiaThang vì TieuChiChungDanhGia không có cong_chuc_id)
    # CHÍNH SÁCH: Spec bắt buộc data phải đã được phê duyệt. Chỉ dùng diem_phe_duyet
    # của các tiêu chí có trang_thai == DA_PHE_DUYET, KHÔNG fallback về tự chấm.
    tieu_chi_list = []
    diem_tieu_chi = 0
    if danh_gia:
        stmt_tc = (
            select(TieuChiChungDanhGia)
            .where(TieuChiChungDanhGia.danh_gia_thang_id == danh_gia.id)
            .where(TieuChiChungDanhGia.trang_thai == TrangThaiTieuChi.DA_PHE_DUYET)
            .options(selectinload(TieuChiChungDanhGia.tieu_chi))
        )
        result_tc = await db.execute(stmt_tc)
        tieu_chi_list = result_tc.scalars().all()

        # Tính điểm tiêu chí chung — chỉ dùng diem_phe_duyet (đã lọc DA_PHE_DUYET)
        diem_tieu_chi = sum([
            float(tc.diem_phe_duyet or 0)
            for tc in tieu_chi_list
        ])

    # Chọn template
    is_lanh_dao = cc.is_lanh_dao or False
    template_name = "PL-Mẫu số 01B-LĐ.docx" if is_lanh_dao else "PL-Mẫu số 01A-CC.docx"
    template_path = TEMPLATES_DIR / template_name

    if not template_path.exists():
        raise HTTPException(status_code=500, detail=f"Template {template_name} không tồn tại")

    # Load template
    doc = Document(template_path)

    # Replace placeholders
    ky_danh_gia = f"Tháng {thang}/{nam}"
    ho_ten = cc.ho_ten or "N/A"
    chuc_vu = cc.chuc_vu or "Công chức"
    don_vi = cc.don_vi.ten_don_vi if cc.don_vi else "N/A"
    diem_tieu_chi_str = f"{diem_tieu_chi:.2f}"

    # Tính điểm tổng (nếu có danh_gia)
    diem_tong_str = "Chưa có dữ liệu"
    if danh_gia and danh_gia.diem_tong is not None:
        diem_tong_str = f"{float(danh_gia.diem_tong):.2f}"

    replace_placeholder_in_docx(doc, "{{ky_danh_gia}}", ky_danh_gia)
    replace_placeholder_in_docx(doc, "{{ho_ten}}", ho_ten)
    replace_placeholder_in_docx(doc, "{{chuc_vu}}", chuc_vu)
    replace_placeholder_in_docx(doc, "{{don_vi}}", don_vi)
    replace_placeholder_in_docx(doc, "{{diem_tieu_chi}}", diem_tieu_chi_str)
    replace_placeholder_in_docx(doc, "{{diem_tong}}", diem_tong_str)

    # === Tính điểm nhiệm vụ REUSE logic từ xep_loai_moi.py ===
    # Import hàm tính điểm KPI 70 (production logic)
    from app.api.v1.endpoints.xep_loai_moi import (
        tinh_diem_kpi_70,
        tinh_diem_kpi_70_lanh_dao,
        tinh_diem_kpi_70_hd_111,
        _has_ke_khai_lanh_dao,
    )

    # Gọi hàm tính điểm (tam_tinh=False vì cần kết quả chính thức)
    diem_a = None
    diem_b = None
    diem_c = None
    diem_d = None
    diem_dd = None
    diem_e = None
    diem_kpi_nhiem_vu = None

    try:
        if is_lanh_dao:
            # Lãnh đạo: dùng công thức 6 chỉ số (a,b,c,d,đ,e)
            kpi_data = await tinh_diem_kpi_70_lanh_dao(db, cc.id, thang, nam, tam_tinh=False)
            # Cap tỷ lệ % ở 100% (phục vụ hiển thị trong phiếu đánh giá)
            diem_a = min(kpi_data.get("a_so_luong", 0) * 100, 100.0)  # a: tỷ lệ hoàn thành
            diem_b = min(kpi_data.get("b_tien_do", 0) * 100, 100.0)   # b: tỷ lệ tiến độ
            diem_c = min(kpi_data.get("c_chat_luong", 0) * 100, 100.0)  # c: tỷ lệ chất lượng
            diem_d = min(kpi_data.get("d_ket_qua", 0) * 100, 100.0)   # d: kết quả đơn vị
            diem_dd = min(kpi_data.get("dd_to_chuc", 0) * 100, 100.0)  # đ: tổ chức triển khai
            diem_e = min(kpi_data.get("e_doan_ket", 0) * 100, 100.0)   # e: đoàn kết nội bộ
            diem_kpi_nhiem_vu = kpi_data.get("diem_70", 0)  # Điểm KPI /70
        elif cc.is_hd_111 and await _has_ke_khai_lanh_dao(db, cc.id, thang, nam):
            # HĐ 111 (Phase 3 — 29/04/2026): công thức 3 chỉ số (a,b,c) lấy từ ke_khai_lanh_dao
            kpi_data = await tinh_diem_kpi_70_hd_111(db, cc.id, thang, nam, tam_tinh=False)
            diem_a = min(kpi_data.get("a_so_luong", 0) * 100, 100.0)
            diem_b = min(kpi_data.get("b_tien_do", 0) * 100, 100.0)
            diem_c = min(kpi_data.get("c_chat_luong", 0) * 100, 100.0)
            diem_kpi_nhiem_vu = kpi_data.get("diem_70", 0)
        else:
            # Công chức: dùng công thức 3 chỉ số (a,b,c).
            # Áp cho cả HĐ 111 ở các tháng cũ (chưa có data ke_khai_lanh_dao).
            kpi_data = await tinh_diem_kpi_70(db, cc.id, thang, nam, tam_tinh=False)
            # Cap tỷ lệ % ở 100% (fix bug a > 100% khi CC làm vượt chỉ tiêu)
            diem_a = min(kpi_data.get("a_so_luong", 0) * 100, 100.0)
            diem_b = min(kpi_data.get("c_tien_do", 0) * 100, 100.0)   # Chú ý: key là "c_tien_do" (theo code)
            diem_c = min(kpi_data.get("b_chat_luong", 0) * 100, 100.0)  # Chú ý: key là "b_chat_luong"
            diem_kpi_nhiem_vu = kpi_data.get("diem_70", 0)
    except Exception as e:
        logger.warning(f"Lỗi tính điểm KPI cho {cc.ma_cc} tháng {thang}/{nam}: {e}")
        # Giữ giá trị None để hiển thị "N/A" trong template

    # Replace vào paragraph (dùng search text)
    # Phải replace toàn bộ text của para rồi gán lại vào run đầu tiên để preserve formatting
    for para in doc.paragraphs:
        full_text = para.text
        replaced = False

        # Điểm a (số lượng)
        # Pattern: "………..(%)" (3× U+2026 + ".." + "(%)")
        if "Điểm tỷ lệ % đánh giá về số lượng (a):" in full_text:
            if diem_a is not None:
                full_text = full_text.replace("………..(%)", f"{diem_a:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm b (tiến độ)
        elif "Điểm tỷ lệ % đánh giá về tiến độ (b):" in full_text:
            if diem_b is not None:
                full_text = full_text.replace("………..(%)", f"{diem_b:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm c (chất lượng)
        elif "Điểm tỷ lệ % đánh giá về chất lượng (c):" in full_text:
            if diem_c is not None:
                full_text = full_text.replace("………..(%)", f"{diem_c:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm d (kết quả đơn vị)
        elif "Điểm tỷ lệ % đánh giá về kết quả hoạt động của bộ phận/đơn vị (d):" in full_text:
            if diem_d is not None:
                full_text = full_text.replace("………..(%)", f"{diem_d:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm đ (tổ chức triển khai) — pattern khác: "……(%)"
        elif "Điểm tỷ lệ % đánh giá về khả năng tổ chức triển khai thực hiện nhiệm vụ (đ):" in full_text:
            if diem_dd is not None:
                full_text = full_text.replace("……(%)", f"{diem_dd:.2f}%")
            else:
                full_text = full_text.replace("……(%)", "N/A")
            replaced = True

        # Điểm e (đoàn kết) — text khác, chứa "năng lực tập hợp"
        elif "Điểm tỷ lệ % đánh giá về năng lực tập hợp, đoàn kết công chức" in full_text:
            if diem_e is not None:
                full_text = full_text.replace("………..(%)", f"{diem_e:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm d (tổng kết quả thực hiện nhiệm vụ - PL-01A cho CC thường)
        # d = (a + b + c) / 3 × 100 (%), cap ở 100%
        # Placeholder: "d. Điểm tỷ lệ % kết quả thực hiện nhiệm vụ (d = ): ……….(%)"
        # FIX BUG: Dùng replace_in_runs() để giữ nguyên equation trong template
        # KHÔNG set replaced=True vì replace_in_runs() đã xử lý xong, không cần clear all runs
        elif "Điểm tỷ lệ % kết quả thực hiện nhiệm vụ (d = )" in full_text:
            diem_d_01a = None
            if not is_lanh_dao and None not in [diem_a, diem_b, diem_c]:
                # CC thường (PL-01A): d = (a + b + c) / 3
                diem_d_01a = min((diem_a + diem_b + diem_c) / 3, 100.0)

            if diem_d_01a is not None:
                replace_in_runs(para, "……….(%)", f"{diem_d_01a:.2f}%")
            else:
                replace_in_runs(para, "……….(%)", "N/A")

        # Điểm g (tổng kết quả thực hiện nhiệm vụ - PL-01B cho lãnh đạo)
        # g = (a + b + c + d + đ + e) / 6 × 100 (%)
        # Hoặc tương đương: g = diem_kpi_nhiem_vu / 70 × 100
        # FIX BUG: Dùng replace_in_runs() để giữ nguyên equation trong template
        # KHÔNG set replaced=True vì replace_in_runs() đã xử lý xong, không cần clear all runs
        elif "Điểm tỷ lệ % kết quả thực hiện nhiệm vụ (g = " in full_text:
            diem_g = None
            if is_lanh_dao and None not in [diem_a, diem_b, diem_c, diem_d, diem_dd, diem_e]:
                # Lãnh đạo: g = (a + b + c + d + đ + e) / 6 (đã scale 0-100, nên lấy trung bình)
                diem_g = min((diem_a + diem_b + diem_c + diem_d + diem_dd + diem_e) / 6, 100.0)
            elif not is_lanh_dao and None not in [diem_a, diem_b, diem_c]:
                # CC thường: g = (a + b + c) / 3 (fallback nếu template dùng g thay vì d)
                diem_g = min((diem_a + diem_b + diem_c) / 3, 100.0)

            if diem_g is not None:
                replace_in_runs(para, "……….(%)", f"{diem_g:.2f}%")
            else:
                replace_in_runs(para, "……….(%)", "N/A")

        # Điểm KPI nhiệm vụ (70 điểm max)
        # FIX BUG: Dùng replace_in_runs() để giữ nguyên equation trong template
        # KHÔNG set replaced=True vì replace_in_runs() đã xử lý xong, không cần clear all runs
        elif "Điểm tiêu chí kết quả thực hiện nhiệm vụ (= )" in full_text:
            if diem_kpi_nhiem_vu is not None:
                replace_in_runs(para, "………………………….", f"{diem_kpi_nhiem_vu:.2f}")
            else:
                replace_in_runs(para, "………………………….", "Chưa có dữ liệu")

        # Tổng điểm theo dõi, đánh giá (3=1+2)
        # 1 = điểm tiêu chí chung (/30)
        # 2 = điểm kết quả thực hiện nhiệm vụ (/70)
        # 3 = 1 + 2 = tổng /100
        # Placeholder là ": \t" (tab character)
        elif "3. Tổng điểm theo dõi, đánh giá (3=1+2):" in full_text:
            if diem_kpi_nhiem_vu is not None and diem_tieu_chi is not None:
                tong_diem = diem_tieu_chi + diem_kpi_nhiem_vu
                # Replace ": \t" → ": {số}"
                full_text = full_text.replace(": \t", f": {tong_diem:.2f}")
            else:
                full_text = full_text.replace(": \t", ": Chưa có dữ liệu")
            replaced = True

        # Nếu có replace, gán lại text và set font
        if replaced:
            # Clear all runs
            for run in para.runs:
                run.text = ""
            # Gán text mới vào run đầu tiên (hoặc tạo mới nếu không có)
            if para.runs:
                para.runs[0].text = full_text
                set_times_new_roman(para.runs[0])
            else:
                new_run = para.add_run(full_text)
                set_times_new_roman(new_run)

    # Điền điểm vào bảng tiêu chí (Table 1 - bảng thứ 2)
    if len(doc.tables) >= 2:
        table_tc = doc.tables[1]
        # Map tieu_chi_list theo ma_tieu_chi (1.1, 1.2, 2.1, ...)
        tc_map = {tc.tieu_chi.ma_tieu_chi: tc for tc in tieu_chi_list if tc.tieu_chi}

        # Duyệt rows (bỏ qua header 2 dòng đầu)
        # Row 0,1: header, Row 2: I (nhóm 1), Row 3: 1.1, Row 4: 1.2, Row 5: II (nhóm 2), ...
        # Logic map: đọc cột TT, match với danh sách 10 tiêu chí:
        # Nhóm I: 1.1, 1.2 (2 tiêu chí)
        # Nhóm II: 2.1, 2.2, 2.3, 2.4 (4 tiêu chí)
        # Nhóm III: 3.1, 3.2, 3.3, 3.4 (4 tiêu chí)
        ma_tieu_chi_list = [
            ("1", "1.1"),
            ("2", "1.2"),
            ("1", "2.1"),
            ("2", "2.2"),
            ("3", "2.3"),
            ("4", "2.4"),
            ("1", "3.1"),
            ("2", "3.2"),
            ("3", "3.3"),
            ("4", "3.4"),
        ]

        # Track thứ tự tiêu chí và tổng nhóm
        tc_idx = 0
        tong_nhom_1 = Decimal(0)
        tong_nhom_2 = Decimal(0)
        tong_nhom_3 = Decimal(0)

        for row_idx, row in enumerate(table_tc.rows[2:], start=0):
            if len(row.cells) >= 4:
                cell_tt = row.cells[0].text.strip()

                # Điền tổng nhóm vào row header (I, II, III)
                # Row 2 (idx=0): I → tổng nhóm 1
                # Row 5 (idx=3): II → tổng nhóm 2
                # Row 10 (idx=8): III → tổng nhóm 3
                if cell_tt == "I":
                    # Tính tổng nhóm 1 sau khi đã duyệt 2 tiêu chí (1.1, 1.2)
                    # Nhưng hiện tại chưa duyệt → sẽ cập nhật sau
                    pass
                elif cell_tt == "II":
                    # Cập nhật tổng nhóm 1 vào row I
                    table_tc.rows[2].cells[3].text = f"{float(tong_nhom_1):.2f}"
                    for para in table_tc.rows[2].cells[3].paragraphs:
                        for run in para.runs:
                            set_times_new_roman(run)
                elif cell_tt == "III":
                    # Cập nhật tổng nhóm 2 vào row II
                    table_tc.rows[5].cells[3].text = f"{float(tong_nhom_2):.2f}"
                    for para in table_tc.rows[5].cells[3].paragraphs:
                        for run in para.runs:
                            set_times_new_roman(run)
                elif cell_tt == "Tổng cộng":
                    # Cập nhật tổng nhóm 3 vào row III
                    table_tc.rows[10].cells[3].text = f"{float(tong_nhom_3):.2f}"
                    for para in table_tc.rows[10].cells[3].paragraphs:
                        for run in para.runs:
                            set_times_new_roman(run)

                    # Cập nhật tổng /30 vào row Tổng cộng
                    tong_30 = tong_nhom_1 + tong_nhom_2 + tong_nhom_3
                    row.cells[3].text = f"{float(tong_30):.2f}"
                    for para in row.cells[3].paragraphs:
                        for run in para.runs:
                            set_times_new_roman(run)
                    continue

                # Match theo TT number để điền điểm tiêu chí con
                if tc_idx < len(ma_tieu_chi_list):
                    expected_tt, ma_tc = ma_tieu_chi_list[tc_idx]
                    # Kiểm tra cell_tt có match không (cột TT chứa "1", "2", "3", ...)
                    if cell_tt == expected_tt:
                        # Tìm điểm trong tc_map (đã lọc DA_PHE_DUYET ở trên)
                        diem = Decimal(0)
                        if ma_tc in tc_map:
                            tc_data = tc_map[ma_tc]
                            # Chỉ dùng điểm phê duyệt chính thức
                            diem = tc_data.diem_phe_duyet or Decimal(0)
                            row.cells[3].text = f"{float(diem):.2f}"
                        else:
                            # Tiêu chí chưa được phê duyệt hoặc chưa có dữ liệu
                            row.cells[3].text = ""

                        # Set font Times New Roman
                        for para in row.cells[3].paragraphs:
                            for run in para.runs:
                                set_times_new_roman(run)

                        # Cộng vào tổng nhóm tương ứng
                        if ma_tc in ["1.1", "1.2"]:
                            tong_nhom_1 += diem
                        elif ma_tc in ["2.1", "2.2", "2.3", "2.4"]:
                            tong_nhom_2 += diem
                        elif ma_tc in ["3.1", "3.2", "3.3", "3.4"]:
                            tong_nhom_3 += diem

                        tc_idx += 1

    # Load phiếu THÁNG từ DB (nếu có) và fill mục 4/5/6 + ngày ký + tên LĐ
    phieu_thang = await _lay_phieu_thang(db, cc.id, thang, nam)
    ngay_ky: Optional[datetime] = None
    ten_nguoi_duyet_override: Optional[str] = None

    if phieu_thang is not None:
        if phieu_thang.uu_diem:
            _fill_section_after_heading(doc, "4. Ưu điểm", phieu_thang.uu_diem)
        if phieu_thang.han_che:
            _fill_section_after_heading(doc, "5. Hạn chế", phieu_thang.han_che)
        if (
            phieu_thang.y_kien_lanh_dao
            and phieu_thang.trang_thai == TrangThaiPhieuDanhGia.DA_PHE_DUYET.value
        ):
            _fill_section_after_heading(doc, "6. Ý kiến", phieu_thang.y_kien_lanh_dao)

        if phieu_thang.trang_thai == TrangThaiPhieuDanhGia.DA_PHE_DUYET.value:
            ngay_ky = phieu_thang.ngay_phe_duyet
        elif phieu_thang.trang_thai == TrangThaiPhieuDanhGia.CHO_PHE_DUYET.value:
            ngay_ky = phieu_thang.ngay_gui_duyet

        if phieu_thang.trang_thai in (
            TrangThaiPhieuDanhGia.DA_PHE_DUYET.value,
            TrangThaiPhieuDanhGia.BI_TU_CHOI.value,
        ) and phieu_thang.nguoi_phe_duyet is not None:
            ten_nguoi_duyet_override = phieu_thang.nguoi_phe_duyet.ho_ten or ""

    # Auto-fill footer: ngày + tên CC + tên TDV/CCT
    await _apply_auto_fill_chung(
        doc, db, cc,
        ngay_ky=ngay_ky,
        ten_nguoi_duyet_override=ten_nguoi_duyet_override,
    )

    # Lưu vào buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Tên file
    ma_cc_safe = cc.ma_cc.replace("/", "-") if cc.ma_cc else "user"
    filename = f"PhieuDanhGia_{ma_cc_safe}_T{thang}_{nam}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# =============================================================================
# ENDPOINT 2: IN BẢNG KÊ CÔNG VIỆC (PL-02)
# =============================================================================

@router.get("/bang-ke-cong-viec/{thang}/{nam}")
async def export_bang_ke_cong_viec(
    thang: int,
    nam: int,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """
    Xuất bảng kê công việc cá nhân (PL-02) cho tháng/năm chỉ định.

    Bao gồm:
    - Thống kê tổng hợp (tổng CV, đã hoàn thành, chưa hoàn thành, ...)
    - Chi tiết từng công việc trong bảng

    Args:
        thang: Tháng (1-12)
        nam: Năm (>= 2020)
        db: Database session
        current_user: User hiện tại

    Returns:
        StreamingResponse: File DOCX download
    """
    # Validation
    if not (1 <= thang <= 12):
        raise HTTPException(status_code=400, detail="Tháng không hợp lệ (1-12)")
    if nam < 2020 or nam > 2100:
        raise HTTPException(status_code=400, detail="Năm không hợp lệ")

    # Load user (kèm vai_tro để check is_hd_111)
    stmt_user = (
        select(CongChuc)
        .options(
            selectinload(CongChuc.don_vi),
            selectinload(CongChuc.vai_tro),
        )
        .where(CongChuc.id == current_user.id)
    )
    result_user = await db.execute(stmt_user)
    cc = result_user.scalar_one_or_none()
    if not cc:
        raise HTTPException(status_code=404, detail="Không tìm thấy thông tin công chức")

    is_lanh_dao = cc.is_lanh_dao or False
    # HĐ 111 (Phase 3 — 29/04/2026) kê khai dùng form Lãnh đạo (ke_khai_lanh_dao),
    # nên PL-02 cũng phải đọc từ KeKhaiLanhDao thay vì KeKhaiCongViec — nếu không
    # bảng kê sẽ rỗng.
    dung_form_lanh_dao = bool(is_lanh_dao or cc.is_hd_111)

    # Lấy danh sách công việc đã kê khai trong tháng
    cong_viec_list = []
    cong_viec_lanh_dao_list = []

    # CHÍNH SÁCH (17/04/2026): Bảng kê chính thức chỉ tổng hợp công việc ĐÃ PHÊ DUYỆT.
    if dung_form_lanh_dao:
        stmt_kkld = (
            select(KeKhaiLanhDao)
            .where(
                and_(
                    KeKhaiLanhDao.cong_chuc_id == cc.id,
                    KeKhaiLanhDao.thang == thang,
                    KeKhaiLanhDao.nam == nam,
                    KeKhaiLanhDao.is_deleted == False,
                    KeKhaiLanhDao.trang_thai == TrangThaiKeKhaiLD.DA_PHE_DUYET.value,
                )
            )
            .order_by(KeKhaiLanhDao.ngay_thuc_hien.desc())
        )
        result_kkld = await db.execute(stmt_kkld)
        cong_viec_lanh_dao_list = result_kkld.scalars().all()
    else:
        stmt_kk = (
            select(KeKhaiCongViec)
            .options(
                selectinload(KeKhaiCongViec.danh_muc_sp).selectinload(DanhMucSpCongViec.sp_chuan),
                selectinload(KeKhaiCongViec.cap_do)
            )
            .where(
                and_(
                    KeKhaiCongViec.cong_chuc_id == cc.id,
                    KeKhaiCongViec.thang == thang,
                    KeKhaiCongViec.nam == nam,
                    KeKhaiCongViec.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET,
                )
            )
            .order_by(KeKhaiCongViec.ngay_thuc_hien.desc())
        )
        result_kk = await db.execute(stmt_kk)
        cong_viec_list = result_kk.scalars().all()

    # Tính thống kê
    if dung_form_lanh_dao:
        # KeKhaiLanhDao: có trang_thai_hoan_thanh (DA_HOAN_THANH/CHUA_HOAN_THANH)
        tong_cv = len(cong_viec_lanh_dao_list)
        da_ht = sum(1 for cv in cong_viec_lanh_dao_list
                    if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH)
        chua_ht = tong_cv - da_ht

        # Tính đúng hạn/trễ hạn cho lãnh đạo dựa vào so_loi_tien_do
        # Đạt tiến độ = đã hoàn thành VÀ không có lỗi tiến độ
        dung_han = sum(
            1 for cv in cong_viec_lanh_dao_list
            if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
            and cv.so_loi_tien_do == 0
        )
        # Chậm tiến độ = đã hoàn thành NHƯNG có lỗi tiến độ
        tre_han = sum(
            1 for cv in cong_viec_lanh_dao_list
            if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
            and cv.so_loi_tien_do > 0
        )

        ty_le_dung_han = round((dung_han / da_ht * 100) if da_ht > 0 else 0, 2)
        ty_le_tre_han = round((tre_han / da_ht * 100) if da_ht > 0 else 0, 2)
    else:
        # KeKhaiCongViec: không có trang_thai_hoan_thanh - chỉ có trang_thai
        tong_cv = len(cong_viec_list)
        # Coi DA_PHE_DUYET là đã hoàn thành
        da_ht = sum(1 for cv in cong_viec_list if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET)
        chua_ht = tong_cv - da_ht

        # Logic mới (fix bug PL-02): Dùng so_loi_tien_do (giống xep_loai_moi.py::tinh_diem_kpi_70)
        # Đạt tiến độ = đã hoàn thành VÀ không có lỗi tiến độ
        dung_han = sum(
            1 for cv in cong_viec_list
            if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
            and cv.so_loi_tien_do == 0
        )
        # Chậm tiến độ = đã hoàn thành NHƯNG có lỗi tiến độ
        tre_han = sum(
            1 for cv in cong_viec_list
            if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
            and cv.so_loi_tien_do > 0
        )

        ty_le_dung_han = round((dung_han / da_ht * 100) if da_ht > 0 else 0, 2)
        ty_le_tre_han = round((tre_han / da_ht * 100) if da_ht > 0 else 0, 2)

    # Load template
    template_path = TEMPLATES_DIR / "PL-Mẫu số 02-Bang ke.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="Template PL-02 không tồn tại")

    doc = Document(template_path)

    # Replace placeholders
    ky_danh_gia = f"Tháng {thang}/{nam}"
    replace_placeholder_in_docx(doc, "{{ky_danh_gia}}", ky_danh_gia)
    replace_placeholder_in_docx(doc, "{{tong_cv}}", str(tong_cv))
    replace_placeholder_in_docx(doc, "{{chua_ht}}", str(chua_ht))
    replace_placeholder_in_docx(doc, "{{da_ht}}", str(da_ht))
    replace_placeholder_in_docx(doc, "{{dung_han}}", str(dung_han))
    replace_placeholder_in_docx(doc, "{{ty_le_dung_han}}", f"{ty_le_dung_han}")
    replace_placeholder_in_docx(doc, "{{tre_han}}", str(tre_han))
    replace_placeholder_in_docx(doc, "{{ty_le_tre_han}}", f"{ty_le_tre_han}")

    # Điền bảng công việc (Table 1 - bảng thứ 2)
    if len(doc.tables) >= 2:
        table_cv = doc.tables[1]

        # Xóa các row mẫu (giữ header 3 dòng đầu)
        # Rows từ index 3 trở đi là data rows mẫu
        rows_to_delete = []
        for idx in range(3, len(table_cv.rows)):
            rows_to_delete.append(idx)

        # Delete từ cuối lên đầu
        for idx in reversed(rows_to_delete):
            table_cv._element.remove(table_cv.rows[idx]._element)

        # Thêm row mới cho từng công việc
        # Chọn danh sách dựa vào kiểu kê khai (LĐ thật + HĐ 111 dùng KeKhaiLanhDao).
        data_source = cong_viec_lanh_dao_list if dung_form_lanh_dao else cong_viec_list

        for idx, cv in enumerate(data_source, start=1):
            # Add new row (copy format từ row template cuối cùng)
            new_row = table_cv.add_row()

            # Điền dữ liệu vào các cột
            # (1) STT
            new_row.cells[0].text = str(idx)
            set_cell_font_times_new_roman(new_row.cells[0])

            if dung_form_lanh_dao:
                # ===== MAPPING CHO KÊ KHAI DẠNG LÃNH ĐẠO (KeKhaiLanhDao) =====
                # (2) Nhiệm vụ được giao (tên công việc)
                new_row.cells[1].text = cv.ten_cong_viec or "N/A"
                set_cell_font_times_new_roman(new_row.cells[1])

                # (3) Số, ngày văn bản giao (không có trong KeKhaiLanhDao - bỏ trống)
                new_row.cells[2].text = ""
                set_cell_font_times_new_roman(new_row.cells[2])

                # (4) Sản phẩm đầu ra — lấy từ KeKhaiLanhDao.mo_ta
                new_row.cells[3].text = cv.mo_ta or ""
                set_cell_font_times_new_roman(new_row.cells[3])

                # (5) Ngày phải hoàn thành (dùng ngay_thuc_hien - giống cột 6)
                ngay_ht_str = cv.ngay_thuc_hien.strftime("%d/%m/%Y") if cv.ngay_thuc_hien else ""
                new_row.cells[4].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[4])

                # (6) Ngày hoàn thành thực tế (dùng ngay_thuc_hien)
                new_row.cells[5].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[5])

                # (7) Chưa hoàn thành
                chua_hoan_thanh = cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.CHUA_HOAN_THANH
                new_row.cells[6].text = "X" if chua_hoan_thanh else ""
                set_cell_font_times_new_roman(new_row.cells[6])

                # (8) Đạt tiến độ (đã hoàn thành và không có lỗi tiến độ)
                dat_tien_do = (
                    cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
                    and cv.so_loi_tien_do == 0
                )
                new_row.cells[7].text = "X" if dat_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[7])

                # (9) Chậm tiến độ (đã hoàn thành nhưng có lỗi tiến độ)
                cham_tien_do = (
                    cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
                    and cv.so_loi_tien_do > 0
                )
                new_row.cells[8].text = "X" if cham_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[8])

                # (10) Ghi chú (gộp lỗi CL + TĐ + số lỗi)
                ghi_chu = ""
                if cv.so_loi_chat_luong > 0:
                    ghi_chu += f"Lỗi CL ({cv.so_loi_chat_luong}): {cv.ghi_chu_loi_chat_luong or 'Không có mô tả'}. "
                if cv.so_loi_tien_do > 0:
                    ghi_chu += f"Lỗi TĐ ({cv.so_loi_tien_do}): {cv.ghi_chu_loi_tien_do or 'Không có mô tả'}."
                if cv.y_kien_lanh_dao:
                    ghi_chu += f" Ý kiến LĐ: {cv.y_kien_lanh_dao}."
                new_row.cells[9].text = ghi_chu.strip()
                set_cell_font_times_new_roman(new_row.cells[9])

            else:
                # ===== MAPPING CHO CÔNG CHỨC THƯỜNG (KeKhaiCongViec) =====
                # (2) Nhiệm vụ được giao (tên công việc + cấp độ)
                ten_sp = ""
                if cv.danh_muc_sp:
                    ten_sp = cv.danh_muc_sp.ten_cong_viec
                    if cv.cap_do:
                        ten_sp += f" ({cv.cap_do.ten_cap_do})"
                new_row.cells[1].text = ten_sp or "N/A"
                set_cell_font_times_new_roman(new_row.cells[1])

                # (3) Số, ngày văn bản giao (không có trong model - bỏ trống)
                new_row.cells[2].text = ""
                set_cell_font_times_new_roman(new_row.cells[2])

                # (4) Sản phẩm đầu ra — lấy từ mo_ta_cong_viec CC tự nhập
                # (rỗng nếu CC bỏ trống — theo policy 2026-05-12).
                new_row.cells[3].text = cv.mo_ta_cong_viec or ""
                set_cell_font_times_new_roman(new_row.cells[3])

                # (5) + (6) Ngày hoàn thành: Dùng ngay_thuc_hien (CC không điền ngay_hoan_thanh)
                # Production data: 152/154 records có ngay_hoan_thanh=NULL, nhưng 100% có ngay_thuc_hien
                ngay_ht_str = cv.ngay_thuc_hien.strftime("%d/%m/%Y") if cv.ngay_thuc_hien else ""
                new_row.cells[4].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[4])

                # (6) Ngày hoàn thành thực tế (dùng ngay_thuc_hien)
                new_row.cells[5].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[5])

                # (7) Chưa hoàn thành
                chua_hoan_thanh = cv.trang_thai != TrangThaiKeKhai.DA_PHE_DUYET
                new_row.cells[6].text = "X" if chua_hoan_thanh else ""
                set_cell_font_times_new_roman(new_row.cells[6])

                # (8) Đạt tiến độ: Dùng logic từ xep_loai_moi.py::tinh_diem_kpi_70
                # Đã hoàn thành VÀ không có lỗi tiến độ (so_loi_tien_do=0)
                dat_tien_do = (
                    cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
                    and cv.so_loi_tien_do == 0
                )
                new_row.cells[7].text = "X" if dat_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[7])

                # (9) Chậm tiến độ: Đã hoàn thành NHƯNG có lỗi tiến độ (so_loi_tien_do>0)
                cham_tien_do = (
                    cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
                    and cv.so_loi_tien_do > 0
                )
                new_row.cells[8].text = "X" if cham_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[8])

                # (10) Ghi chú (gộp lỗi CL + TĐ)
                ghi_chu = ""
                if cv.ghi_chu_loi_chat_luong:
                    ghi_chu += f"Lỗi CL: {cv.ghi_chu_loi_chat_luong}. "
                if cv.ghi_chu_loi_tien_do:
                    ghi_chu += f"Lỗi TĐ: {cv.ghi_chu_loi_tien_do}."
                new_row.cells[9].text = ghi_chu.strip()
                set_cell_font_times_new_roman(new_row.cells[9])

    # Auto-fill footer (Quảng Ninh + ngày + tên NGƯỜI TỔNG HỢP + THỦ TRƯỞNG ĐƠN VỊ)
    await _apply_auto_fill_chung(doc, db, cc, ngay_ky=None)

    # Lưu vào buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Tên file
    ma_cc_safe = cc.ma_cc.replace("/", "-") if cc.ma_cc else "user"
    filename = f"BangKeCongViec_{ma_cc_safe}_T{thang}_{nam}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# =============================================================================
# ENDPOINT 3: IN PHIẾU ĐÁNH GIÁ QUÝ (PL-01A hoặc PL-01B)
# =============================================================================

@router.get("/phieu-danh-gia-quy/{quy}/{nam}")
async def export_phieu_danh_gia_quy(
    quy: int,
    nam: int,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """
    Xuất phiếu đánh giá cá nhân QUÝ (PL-01A hoặc PL-01B).

    Tương tự endpoint tháng, nhưng:
    - Dùng `tinh_diem_quy()` từ xep_loai_quy_helpers
    - Kỳ đánh giá: "Quý {quy}/{nam}"
    - Điểm a,b,c: trung bình 3 tháng (CC) hoặc a-e theo logic LD (từ tinh_diem_quy)
    - Điểm tiêu chí chung: không điền chi tiết từng ô (deduplicate), chỉ điền tổng

    Args:
        quy: Quý (1-4)
        nam: Năm (>= 2020)
        db: Database session
        current_user: User hiện tại

    Returns:
        StreamingResponse: File DOCX download
    """
    # Import helper
    from app.api.v1.endpoints.xep_loai_quy_helpers import tinh_diem_quy, QUY_TO_THANG

    # Validation
    if not (1 <= quy <= 4):
        raise HTTPException(status_code=400, detail="Quý không hợp lệ (1-4)")
    if nam < 2020 or nam > 2100:
        raise HTTPException(status_code=400, detail="Năm không hợp lệ")

    # Load user
    stmt = (
        select(CongChuc)
        .options(
            selectinload(CongChuc.don_vi),
            selectinload(CongChuc.vai_tro),
        )
        .where(CongChuc.id == current_user.id)
    )
    result = await db.execute(stmt)
    cc = result.scalar_one_or_none()
    if not cc:
        raise HTTPException(status_code=404, detail="Không tìm thấy thông tin công chức")

    is_lanh_dao = cc.is_lanh_dao or False

    # Điểm quý — mặc định lấy `tam_tinh=False` (chỉ DA_PHE_DUYET) cho phiếu chính thức.
    # CHÍNH SÁCH (20/04/2026): với TDV và PCCT — luôn lấy điểm TẠM TÍNH (gộp cả
    # CHO_PHE_DUYET) vì CCT là bottleneck duyệt kê khai của họ; bắt chờ CCT duyệt
    # xong mới in phiếu sẽ làm điểm tụt giả tạo. CCT và CC thường giữ nguyên
    # `tam_tinh=False`.
    cap_bac = cc.vai_tro.cap_bac if cc.vai_tro else None
    dung_tam_tinh = cap_bac in (CapBacVaiTro.TRUONG_DON_VI, CapBacVaiTro.PHO_CHI_CUC_TRUONG)
    ket_qua = await tinh_diem_quy(db, cc.id, quy, nam, tam_tinh=dung_tam_tinh)

    # Lấy 3 tháng trong quý
    thang_list = QUY_TO_THANG.get(quy, [])

    # Điểm tổng hợp — đều lấy trực tiếp từ `tinh_diem_quy` (SINGLE SOURCE OF TRUTH).
    # Công thức lũy kế theo Phụ lục I đã nằm trong hàm đó; ở đây chỉ format cho template.
    diem_tieu_chi = ket_qua.get("diem_tc_quy") or 0
    diem_kpi_nhiem_vu = ket_qua.get("diem_kpi_quy") or 0
    diem_tong = ket_qua.get("diem_tong_quy") or 0

    # a/b/c/d/đ/e quý: 0-1 → đổi sang % cho template.
    # - `tinh_diem_quy` dùng quy ước spec (b=tiến độ, c=chất lượng).
    # - Với CC thường, d/đ/e = None.
    diem_a = (ket_qua.get("a_so_luong") or 0) * 100
    diem_b = (ket_qua.get("b_tien_do") or 0) * 100
    diem_c = (ket_qua.get("c_chat_luong") or 0) * 100
    diem_d = (ket_qua.get("d_ket_qua") * 100) if ket_qua.get("d_ket_qua") is not None else None
    diem_dd = (ket_qua.get("dd_to_chuc") * 100) if ket_qua.get("dd_to_chuc") is not None else None
    diem_e = (ket_qua.get("e_doan_ket") * 100) if ket_qua.get("e_doan_ket") is not None else None

    # Chọn template
    template_name = "PL-Mẫu số 01B-LĐ.docx" if is_lanh_dao else "PL-Mẫu số 01A-CC.docx"
    template_path = TEMPLATES_DIR / template_name

    if not template_path.exists():
        raise HTTPException(status_code=500, detail=f"Template {template_name} không tồn tại")

    # Load template
    doc = Document(template_path)

    # Replace placeholders
    ky_danh_gia = f"Quý {quy}/{nam}"
    ho_ten = cc.ho_ten or "N/A"
    chuc_vu = cc.chuc_vu or "Công chức"
    don_vi = cc.don_vi.ten_don_vi if cc.don_vi else "N/A"
    diem_tieu_chi_str = f"{diem_tieu_chi:.2f}"
    diem_tong_str = f"{diem_tong:.2f}" if diem_tong else "Chưa có dữ liệu"

    replace_placeholder_in_docx(doc, "{{ky_danh_gia}}", ky_danh_gia)
    replace_placeholder_in_docx(doc, "{{ho_ten}}", ho_ten)
    replace_placeholder_in_docx(doc, "{{chuc_vu}}", chuc_vu)
    replace_placeholder_in_docx(doc, "{{don_vi}}", don_vi)
    replace_placeholder_in_docx(doc, "{{diem_tieu_chi}}", diem_tieu_chi_str)
    replace_placeholder_in_docx(doc, "{{diem_tong}}", diem_tong_str)

    # === Replace điểm a,b,c,d,đ,e vào paragraphs ===
    for para in doc.paragraphs:
        full_text = para.text
        replaced = False

        # Điểm a (số lượng)
        if "Điểm tỷ lệ % đánh giá về số lượng (a):" in full_text:
            if diem_a is not None:
                full_text = full_text.replace("………..(%)", f"{diem_a:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm b (tiến độ)
        elif "Điểm tỷ lệ % đánh giá về tiến độ (b):" in full_text:
            if diem_b is not None:
                full_text = full_text.replace("………..(%)", f"{diem_b:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm c (chất lượng)
        elif "Điểm tỷ lệ % đánh giá về chất lượng (c):" in full_text:
            if diem_c is not None:
                full_text = full_text.replace("………..(%)", f"{diem_c:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm d (kết quả đơn vị)
        elif "Điểm tỷ lệ % đánh giá về kết quả hoạt động của bộ phận/đơn vị (d):" in full_text:
            if diem_d is not None:
                full_text = full_text.replace("………..(%)", f"{diem_d:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm đ (tổ chức triển khai)
        elif "Điểm tỷ lệ % đánh giá về khả năng tổ chức triển khai thực hiện nhiệm vụ (đ):" in full_text:
            if diem_dd is not None:
                full_text = full_text.replace("……(%)", f"{diem_dd:.2f}%")
            else:
                full_text = full_text.replace("……(%)", "N/A")
            replaced = True

        # Điểm e (đoàn kết)
        elif "Điểm tỷ lệ % đánh giá về năng lực tập hợp, đoàn kết công chức" in full_text:
            if diem_e is not None:
                full_text = full_text.replace("………..(%)", f"{diem_e:.2f}%")
            else:
                full_text = full_text.replace("………..(%)", "N/A")
            replaced = True

        # Điểm d (tổng kết quả - PL-01A cho CC thường)
        elif "Điểm tỷ lệ % kết quả thực hiện nhiệm vụ (d = )" in full_text:
            diem_d_01a = None
            if not is_lanh_dao and None not in [diem_a, diem_b, diem_c]:
                diem_d_01a = min((diem_a + diem_b + diem_c) / 3, 100.0)

            if diem_d_01a is not None:
                replace_in_runs(para, "……….(%)", f"{diem_d_01a:.2f}%")
            else:
                replace_in_runs(para, "……….(%)", "N/A")

        # Điểm g (tổng kết quả - PL-01B cho lãnh đạo)
        elif "Điểm tỷ lệ % kết quả thực hiện nhiệm vụ (g = " in full_text:
            diem_g = None
            if is_lanh_dao and None not in [diem_a, diem_b, diem_c, diem_d, diem_dd, diem_e]:
                diem_g = min((diem_a + diem_b + diem_c + diem_d + diem_dd + diem_e) / 6, 100.0)
            elif not is_lanh_dao and None not in [diem_a, diem_b, diem_c]:
                diem_g = min((diem_a + diem_b + diem_c) / 3, 100.0)

            if diem_g is not None:
                replace_in_runs(para, "……….(%)", f"{diem_g:.2f}%")
            else:
                replace_in_runs(para, "……….(%)", "N/A")

        # Điểm KPI nhiệm vụ (70 điểm max)
        elif "Điểm tiêu chí kết quả thực hiện nhiệm vụ (= )" in full_text:
            if diem_kpi_nhiem_vu is not None:
                replace_in_runs(para, "………………………….", f"{diem_kpi_nhiem_vu:.2f}")
            else:
                replace_in_runs(para, "………………………….", "Chưa có dữ liệu")

        # Tổng điểm (3=1+2)
        elif "3. Tổng điểm theo dõi, đánh giá (3=1+2):" in full_text:
            if diem_kpi_nhiem_vu is not None and diem_tieu_chi is not None:
                tong_diem_text = diem_tieu_chi + diem_kpi_nhiem_vu
                full_text = full_text.replace(": \t", f": {tong_diem_text:.2f}")
            else:
                full_text = full_text.replace(": \t", ": Chưa có dữ liệu")
            replaced = True

        # Gán lại text nếu có replace
        if replaced:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full_text
                set_times_new_roman(para.runs[0])
            else:
                new_run = para.add_run(full_text)
                set_times_new_roman(new_run)

    # === Điền điểm tiêu chí chung vào bảng (Table 1 - bảng thứ 2) ===
    # Logic quý: Nhóm I, II → MIN (vi phạm bất kỳ tháng nào → vi phạm quý)
    #             Nhóm III → MAX (đạt bất kỳ tháng nào → đạt quý)
    if len(doc.tables) >= 2:
        table_tc = doc.tables[1]

        # Query tiêu chí chung của 3 tháng
        # Lấy tất cả DanhGiaThang trong quý
        stmt_dg_quy = (
            select(DanhGiaThang)
            .where(
                and_(
                    DanhGiaThang.cong_chuc_id == cc.id,
                    DanhGiaThang.thang.in_(thang_list),
                    DanhGiaThang.nam == nam,
                    DanhGiaThang.is_deleted == False,
                )
            )
        )
        result_dg_quy = await db.execute(stmt_dg_quy)
        danh_gia_list = result_dg_quy.scalars().all()
        dg_ids = [dg.id for dg in danh_gia_list]

        # Lấy tất cả TieuChiChungDanhGia của 3 tháng — CHỈ tiêu chí đã phê duyệt.
        # Tháng nào thiếu tiêu chí đã duyệt sẽ tính 0 cho tháng đó (TB /3 vẫn chia 3).
        tc_quy_map = {}  # ma_tieu_chi → list of (diem, nhom, danh_gia_thang_id)
        if dg_ids:
            stmt_tc_quy = (
                select(TieuChiChungDanhGia)
                .where(TieuChiChungDanhGia.danh_gia_thang_id.in_(dg_ids))
                .where(TieuChiChungDanhGia.trang_thai == TrangThaiTieuChi.DA_PHE_DUYET)
                .options(selectinload(TieuChiChungDanhGia.tieu_chi))
            )
            result_tc_quy = await db.execute(stmt_tc_quy)
            all_tc = result_tc_quy.scalars().all()

            for tc in all_tc:
                if tc.tieu_chi:
                    ma = tc.tieu_chi.ma_tieu_chi
                    diem = tc.diem_phe_duyet or Decimal(0)
                    nhom = tc.tieu_chi.nhom_tieu_chi  # 1, 2, 3
                    diem_toi_da = tc.tieu_chi.diem_toi_da or Decimal(0)
                    if ma not in tc_quy_map:
                        tc_quy_map[ma] = {"diem_values": [], "nhom": nhom, "diem_toi_da": diem_toi_da}
                    tc_quy_map[ma]["diem_values"].append(diem)

        # Tính điểm quý cho mỗi tiêu chí: trung bình 3 tháng (luôn chia 3, thiếu tính 0)
        tc_quy_diem = {}  # ma_tieu_chi → Decimal (điểm quý)
        for ma, info in tc_quy_map.items():
            values = info["diem_values"]
            tc_quy_diem[ma] = sum(values) / Decimal("3")

        # Điền vào bảng (logic giống endpoint tháng)
        ma_tieu_chi_list = [
            ("1", "1.1"), ("2", "1.2"),
            ("1", "2.1"), ("2", "2.2"), ("3", "2.3"), ("4", "2.4"),
            ("1", "3.1"), ("2", "3.2"), ("3", "3.3"), ("4", "3.4"),
        ]

        tc_idx = 0
        tong_nhom_1 = Decimal(0)
        tong_nhom_2 = Decimal(0)
        tong_nhom_3 = Decimal(0)

        for row_idx, row in enumerate(table_tc.rows[2:], start=0):
            if len(row.cells) >= 4:
                cell_tt = row.cells[0].text.strip()

                if cell_tt == "I":
                    pass
                elif cell_tt == "II":
                    table_tc.rows[2].cells[3].text = f"{float(tong_nhom_1):.2f}"
                    for p in table_tc.rows[2].cells[3].paragraphs:
                        for r in p.runs:
                            set_times_new_roman(r)
                elif cell_tt == "III":
                    table_tc.rows[5].cells[3].text = f"{float(tong_nhom_2):.2f}"
                    for p in table_tc.rows[5].cells[3].paragraphs:
                        for r in p.runs:
                            set_times_new_roman(r)
                elif cell_tt == "Tổng cộng":
                    table_tc.rows[10].cells[3].text = f"{float(tong_nhom_3):.2f}"
                    for p in table_tc.rows[10].cells[3].paragraphs:
                        for r in p.runs:
                            set_times_new_roman(r)
                    tong_30 = tong_nhom_1 + tong_nhom_2 + tong_nhom_3
                    row.cells[3].text = f"{float(tong_30):.2f}"
                    for p in row.cells[3].paragraphs:
                        for r in p.runs:
                            set_times_new_roman(r)
                    continue

                if tc_idx < len(ma_tieu_chi_list):
                    expected_tt, ma_tc = ma_tieu_chi_list[tc_idx]
                    if cell_tt == expected_tt:
                        diem = Decimal(0)
                        if ma_tc in tc_quy_diem:
                            diem = tc_quy_diem[ma_tc]
                            row.cells[3].text = f"{float(diem):.2f}"
                        else:
                            row.cells[3].text = ""

                        for p in row.cells[3].paragraphs:
                            for r in p.runs:
                                set_times_new_roman(r)

                        if ma_tc in ["1.1", "1.2"]:
                            tong_nhom_1 += diem
                        elif ma_tc in ["2.1", "2.2", "2.3", "2.4"]:
                            tong_nhom_2 += diem
                        elif ma_tc in ["3.1", "3.2", "3.3", "3.4"]:
                            tong_nhom_3 += diem

                        tc_idx += 1

    # === Load phiếu quý từ DB (nếu có) và fill mục 4/5/6 ===
    phieu = await _lay_phieu_quy(db, cc.id, quy, nam)
    ngay_ky: Optional[datetime] = None
    ten_nguoi_duyet_override: Optional[str] = None

    if phieu is not None:
        # Fill nội dung mục 4/5/6 (nếu có)
        if phieu.uu_diem:
            _fill_section_after_heading(doc, "4. Ưu điểm", phieu.uu_diem)
        if phieu.han_che:
            _fill_section_after_heading(doc, "5. Hạn chế", phieu.han_che)
        if phieu.y_kien_lanh_dao and phieu.trang_thai == TrangThaiPhieuDanhGia.DA_PHE_DUYET.value:
            _fill_section_after_heading(doc, "6. Ý kiến", phieu.y_kien_lanh_dao)

        # Ngày ký footer theo trạng thái
        if phieu.trang_thai == TrangThaiPhieuDanhGia.DA_PHE_DUYET.value:
            ngay_ky = phieu.ngay_phe_duyet
        elif phieu.trang_thai == TrangThaiPhieuDanhGia.CHO_PHE_DUYET.value:
            ngay_ky = phieu.ngay_gui_duyet

        # Tên người duyệt:
        # - DA_PHE_DUYET / BI_TU_CHOI: đã có nguoi_phe_duyet → dùng luôn
        # - NHAP / CHO_PHE_DUYET: auto-tra theo rule (để trống nếu không có)
        if phieu.trang_thai in (
            TrangThaiPhieuDanhGia.DA_PHE_DUYET.value,
            TrangThaiPhieuDanhGia.BI_TU_CHOI.value,
        ) and phieu.nguoi_phe_duyet is not None:
            ten_nguoi_duyet_override = phieu.nguoi_phe_duyet.ho_ten or ""

    # Auto-fill footer: ngày + tên CC + tên TDV/CCT
    await _apply_auto_fill_chung(
        doc, db, cc,
        ngay_ky=ngay_ky,
        ten_nguoi_duyet_override=ten_nguoi_duyet_override,
    )

    # Lưu vào buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Tên file
    ma_cc_safe = cc.ma_cc.replace("/", "-") if cc.ma_cc else "user"
    filename = f"PhieuDanhGia_{ma_cc_safe}_Q{quy}_{nam}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# =============================================================================
# ENDPOINT 4: IN BẢNG KÊ CÔNG VIỆC QUÝ (PL-02)
# =============================================================================

@router.get("/bang-ke-cong-viec-quy/{quy}/{nam}")
async def export_bang_ke_cong_viec_quy(
    quy: int,
    nam: int,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """
    Xuất bảng kê công việc cá nhân QUÝ (PL-02).

    Gộp công việc cả 3 tháng trong quý vào 1 bảng.

    Args:
        quy: Quý (1-4)
        nam: Năm (>= 2020)
        db: Database session
        current_user: User hiện tại

    Returns:
        StreamingResponse: File DOCX download
    """
    # Import helper
    from app.api.v1.endpoints.xep_loai_quy_helpers import QUY_TO_THANG

    # Validation
    if not (1 <= quy <= 4):
        raise HTTPException(status_code=400, detail="Quý không hợp lệ (1-4)")
    if nam < 2020 or nam > 2100:
        raise HTTPException(status_code=400, detail="Năm không hợp lệ")

    # Load user (kèm vai_tro để check cap_bac)
    stmt_user = (
        select(CongChuc)
        .options(
            selectinload(CongChuc.don_vi),
            selectinload(CongChuc.vai_tro),
        )
        .where(CongChuc.id == current_user.id)
    )
    result_user = await db.execute(stmt_user)
    cc = result_user.scalar_one_or_none()
    if not cc:
        raise HTTPException(status_code=404, detail="Không tìm thấy thông tin công chức")

    is_lanh_dao = cc.is_lanh_dao or False
    # HĐ 111 (Phase 3 — 29/04/2026) kê khai dùng form Lãnh đạo (ke_khai_lanh_dao),
    # nên PL-02 quý cũng phải đọc từ KeKhaiLanhDao thay vì KeKhaiCongViec.
    dung_form_lanh_dao = bool(is_lanh_dao or cc.is_hd_111)

    # CHÍNH SÁCH (20/04/2026): TDV và PCCT — liệt kê thêm CV CHO_PHE_DUYET (đánh
    # nhãn "[Chờ duyệt]") để đồng bộ với PL-01B quý (dùng tạm tính). CCT và
    # CC thường giữ nguyên chỉ lấy DA_PHE_DUYET.
    cap_bac = cc.vai_tro.cap_bac if cc.vai_tro else None
    gom_cho_duyet = cap_bac in (CapBacVaiTro.TRUONG_DON_VI, CapBacVaiTro.PHO_CHI_CUC_TRUONG)

    # Lấy 3 tháng trong quý
    thang_list = QUY_TO_THANG.get(quy, [])

    # Lấy danh sách công việc trong 3 tháng của quý
    cong_viec_list = []
    cong_viec_lanh_dao_list = []

    # CHÍNH SÁCH (17/04/2026): Bảng kê quý chính thức chỉ tổng hợp công việc ĐÃ PHÊ DUYỆT,
    # ngoại lệ cho TDV/PCCT (xem chú thích gom_cho_duyet ở trên).
    if dung_form_lanh_dao:
        allowed_ld = [TrangThaiKeKhaiLD.DA_PHE_DUYET.value]
        if gom_cho_duyet:
            allowed_ld.append(TrangThaiKeKhaiLD.CHO_PHE_DUYET.value)
        stmt_kkld = (
            select(KeKhaiLanhDao)
            .where(
                and_(
                    KeKhaiLanhDao.cong_chuc_id == cc.id,
                    KeKhaiLanhDao.thang.in_(thang_list),
                    KeKhaiLanhDao.nam == nam,
                    KeKhaiLanhDao.is_deleted == False,
                    KeKhaiLanhDao.trang_thai.in_(allowed_ld),
                )
            )
            .order_by(KeKhaiLanhDao.thang, KeKhaiLanhDao.ngay_thuc_hien.desc())
        )
        result_kkld = await db.execute(stmt_kkld)
        cong_viec_lanh_dao_list = result_kkld.scalars().all()
    else:
        stmt_kk = (
            select(KeKhaiCongViec)
            .options(
                selectinload(KeKhaiCongViec.danh_muc_sp).selectinload(DanhMucSpCongViec.sp_chuan),
                selectinload(KeKhaiCongViec.cap_do)
            )
            .where(
                and_(
                    KeKhaiCongViec.cong_chuc_id == cc.id,
                    KeKhaiCongViec.thang.in_(thang_list),
                    KeKhaiCongViec.nam == nam,
                    KeKhaiCongViec.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET,
                )
            )
            .order_by(KeKhaiCongViec.thang, KeKhaiCongViec.ngay_thuc_hien.desc())
        )
        result_kk = await db.execute(stmt_kk)
        cong_viec_list = result_kk.scalars().all()

    # Tính thống kê (cộng gộp 3 tháng)
    if dung_form_lanh_dao:
        tong_cv = len(cong_viec_lanh_dao_list)
        da_ht = sum(1 for cv in cong_viec_lanh_dao_list
                    if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH)
        chua_ht = tong_cv - da_ht

        dung_han = sum(
            1 for cv in cong_viec_lanh_dao_list
            if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
            and cv.so_loi_tien_do == 0
        )
        tre_han = sum(
            1 for cv in cong_viec_lanh_dao_list
            if cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
            and cv.so_loi_tien_do > 0
        )

        ty_le_dung_han = round((dung_han / da_ht * 100) if da_ht > 0 else 0, 2)
        ty_le_tre_han = round((tre_han / da_ht * 100) if da_ht > 0 else 0, 2)
    else:
        tong_cv = len(cong_viec_list)
        da_ht = sum(1 for cv in cong_viec_list if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET)
        chua_ht = tong_cv - da_ht

        dung_han = sum(
            1 for cv in cong_viec_list
            if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
            and cv.so_loi_tien_do == 0
        )
        tre_han = sum(
            1 for cv in cong_viec_list
            if cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
            and cv.so_loi_tien_do > 0
        )

        ty_le_dung_han = round((dung_han / da_ht * 100) if da_ht > 0 else 0, 2)
        ty_le_tre_han = round((tre_han / da_ht * 100) if da_ht > 0 else 0, 2)

    # Load template
    template_path = TEMPLATES_DIR / "PL-Mẫu số 02-Bang ke.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="Template PL-02 không tồn tại")

    doc = Document(template_path)

    # Replace placeholders
    ky_danh_gia = f"Quý {quy}/{nam}"
    replace_placeholder_in_docx(doc, "{{ky_danh_gia}}", ky_danh_gia)
    replace_placeholder_in_docx(doc, "{{tong_cv}}", str(tong_cv))
    replace_placeholder_in_docx(doc, "{{chua_ht}}", str(chua_ht))
    replace_placeholder_in_docx(doc, "{{da_ht}}", str(da_ht))
    replace_placeholder_in_docx(doc, "{{dung_han}}", str(dung_han))
    replace_placeholder_in_docx(doc, "{{ty_le_dung_han}}", f"{ty_le_dung_han}")
    replace_placeholder_in_docx(doc, "{{tre_han}}", str(tre_han))
    replace_placeholder_in_docx(doc, "{{ty_le_tre_han}}", f"{ty_le_tre_han}")

    # Điền bảng công việc (Table 1 - bảng thứ 2)
    if len(doc.tables) >= 2:
        table_cv = doc.tables[1]

        # Xóa các row mẫu (giữ header 3 dòng đầu)
        rows_to_delete = []
        for idx in range(3, len(table_cv.rows)):
            rows_to_delete.append(idx)

        for idx in reversed(rows_to_delete):
            table_cv._element.remove(table_cv.rows[idx]._element)

        # Thêm row mới cho từng công việc (LĐ thật + HĐ 111 dùng KeKhaiLanhDao)
        data_source = cong_viec_lanh_dao_list if dung_form_lanh_dao else cong_viec_list

        for idx, cv in enumerate(data_source, start=1):
            new_row = table_cv.add_row()

            # (1) STT
            new_row.cells[0].text = str(idx)
            set_cell_font_times_new_roman(new_row.cells[0])

            if dung_form_lanh_dao:
                # ===== MAPPING CHO KÊ KHAI DẠNG LÃNH ĐẠO (LĐ thật + HĐ 111) =====
                new_row.cells[1].text = cv.ten_cong_viec or "N/A"
                set_cell_font_times_new_roman(new_row.cells[1])

                new_row.cells[2].text = ""
                set_cell_font_times_new_roman(new_row.cells[2])

                new_row.cells[3].text = cv.mo_ta or ""
                set_cell_font_times_new_roman(new_row.cells[3])

                ngay_ht_str = cv.ngay_thuc_hien.strftime("%d/%m/%Y") if cv.ngay_thuc_hien else ""
                new_row.cells[4].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[4])

                new_row.cells[5].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[5])

                chua_hoan_thanh = cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.CHUA_HOAN_THANH
                new_row.cells[6].text = "X" if chua_hoan_thanh else ""
                set_cell_font_times_new_roman(new_row.cells[6])

                dat_tien_do = (
                    cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
                    and cv.so_loi_tien_do == 0
                )
                new_row.cells[7].text = "X" if dat_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[7])

                cham_tien_do = (
                    cv.trang_thai_hoan_thanh == TrangThaiHoanThanh.DA_HOAN_THANH
                    and cv.so_loi_tien_do > 0
                )
                new_row.cells[8].text = "X" if cham_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[8])

                ghi_chu_parts = []
                # Nhãn "Chờ duyệt" cho CV chưa được CCT phê duyệt (chỉ xuất hiện với
                # TDV/PCCT — xem `gom_cho_duyet`).
                if cv.trang_thai == TrangThaiKeKhaiLD.CHO_PHE_DUYET.value:
                    ghi_chu_parts.append("[Chờ duyệt]")
                if cv.so_loi_chat_luong > 0:
                    ghi_chu_parts.append(f"Lỗi CL ({cv.so_loi_chat_luong}): {cv.ghi_chu_loi_chat_luong or 'Không có mô tả'}.")
                if cv.so_loi_tien_do > 0:
                    ghi_chu_parts.append(f"Lỗi TĐ ({cv.so_loi_tien_do}): {cv.ghi_chu_loi_tien_do or 'Không có mô tả'}.")
                if cv.y_kien_lanh_dao:
                    ghi_chu_parts.append(f"Ý kiến LĐ: {cv.y_kien_lanh_dao}.")
                new_row.cells[9].text = " ".join(ghi_chu_parts)
                set_cell_font_times_new_roman(new_row.cells[9])

            else:
                # ===== MAPPING CHO CÔNG CHỨC THƯỜNG =====
                ten_sp = ""
                if cv.danh_muc_sp:
                    ten_sp = cv.danh_muc_sp.ten_cong_viec
                    if cv.cap_do:
                        ten_sp += f" ({cv.cap_do.ten_cap_do})"
                new_row.cells[1].text = ten_sp or "N/A"
                set_cell_font_times_new_roman(new_row.cells[1])

                new_row.cells[2].text = ""
                set_cell_font_times_new_roman(new_row.cells[2])

                # (4) Sản phẩm đầu ra — lấy từ mo_ta_cong_viec CC tự nhập
                # (rỗng nếu CC bỏ trống — theo policy 2026-05-12).
                new_row.cells[3].text = cv.mo_ta_cong_viec or ""
                set_cell_font_times_new_roman(new_row.cells[3])

                ngay_ht_str = cv.ngay_thuc_hien.strftime("%d/%m/%Y") if cv.ngay_thuc_hien else ""
                new_row.cells[4].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[4])

                new_row.cells[5].text = ngay_ht_str
                set_cell_font_times_new_roman(new_row.cells[5])

                chua_hoan_thanh = cv.trang_thai != TrangThaiKeKhai.DA_PHE_DUYET
                new_row.cells[6].text = "X" if chua_hoan_thanh else ""
                set_cell_font_times_new_roman(new_row.cells[6])

                dat_tien_do = (
                    cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
                    and cv.so_loi_tien_do == 0
                )
                new_row.cells[7].text = "X" if dat_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[7])

                cham_tien_do = (
                    cv.trang_thai == TrangThaiKeKhai.DA_PHE_DUYET
                    and cv.so_loi_tien_do > 0
                )
                new_row.cells[8].text = "X" if cham_tien_do else ""
                set_cell_font_times_new_roman(new_row.cells[8])

                ghi_chu = ""
                if cv.ghi_chu_loi_chat_luong:
                    ghi_chu += f"Lỗi CL: {cv.ghi_chu_loi_chat_luong}. "
                if cv.ghi_chu_loi_tien_do:
                    ghi_chu += f"Lỗi TĐ: {cv.ghi_chu_loi_tien_do}."
                new_row.cells[9].text = ghi_chu.strip()
                set_cell_font_times_new_roman(new_row.cells[9])

    # Auto-fill footer (Quảng Ninh + ngày + tên NGƯỜI TỔNG HỢP + THỦ TRƯỞNG ĐƠN VỊ)
    await _apply_auto_fill_chung(doc, db, cc, ngay_ky=None)

    # Lưu vào buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Tên file
    ma_cc_safe = cc.ma_cc.replace("/", "-") if cc.ma_cc else "user"
    filename = f"BangKeCongViec_{ma_cc_safe}_Q{quy}_{nam}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# =============================================================================
# ENDPOINTS DÀNH CHO TDV/CCT: TẢI BẢNG IN CỦA CÔNG CHỨC KHÁC
# =============================================================================
# Khi TDV duyệt phiếu quý, có thể cần tải về phiếu / bảng kê bản in của CC
# để đính kèm hồ sơ. 2 endpoint sau reuse lại 2 route quý hiện có bằng cách
# truyền `target_cc` thay cho `current_user` khi gọi hàm xử lý bên trong.


async def _load_cc_cho_in_boi_tdv(
    db,
    nguoi_in: CongChuc,
    target_cong_chuc_id: UUID,
) -> CongChuc:
    """
    Load công chức đích và kiểm tra `nguoi_in` có quyền tải bảng in của CC đó không.

    Quy tắc (đồng bộ với `phieu_danh_gia_quy._co_quyen_duyet`):
    - TDV cùng đơn vị → tải được bảng in của CC / Phó ĐV trong đơn vị mình.
    - CCT → tải được bảng in của TDV / Phó CCT toàn chi cục.
    - Chính CC vẫn được tải bảng in của mình (phục vụ fallback).
    """
    stmt = (
        select(CongChuc)
        .options(
            selectinload(CongChuc.don_vi),
            selectinload(CongChuc.vai_tro),
        )
        .where(CongChuc.id == target_cong_chuc_id)
    )
    target_cc = (await db.execute(stmt)).scalar_one_or_none()
    if target_cc is None:
        raise HTTPException(status_code=404, detail="Không tìm thấy công chức đích")

    if nguoi_in.id == target_cc.id:
        return target_cc

    cb_nguoi = nguoi_in.vai_tro.cap_bac if nguoi_in.vai_tro else None
    cb_target = target_cc.vai_tro.cap_bac if target_cc.vai_tro else None

    allowed = False
    if cb_target in (CapBacVaiTro.CONG_CHUC, CapBacVaiTro.PHO_DON_VI):
        allowed = (
            cb_nguoi == CapBacVaiTro.TRUONG_DON_VI
            and nguoi_in.don_vi_id == target_cc.don_vi_id
        )
    elif cb_target in (CapBacVaiTro.TRUONG_DON_VI, CapBacVaiTro.PHO_CHI_CUC_TRUONG):
        allowed = cb_nguoi == CapBacVaiTro.CHI_CUC_TRUONG

    if not allowed:
        raise HTTPException(
            status_code=403,
            detail="Bạn không có quyền tải bảng in của công chức này",
        )

    return target_cc


@router.get("/phieu-danh-gia-quy/{quy}/{nam}/cua-cc/{cong_chuc_id}")
async def export_phieu_danh_gia_quy_cua_cc(
    quy: int,
    nam: int,
    cong_chuc_id: UUID,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """TDV/CCT tải phiếu đánh giá quý (PL-01A/01B) của CC cụ thể."""
    target_cc = await _load_cc_cho_in_boi_tdv(db, current_user, cong_chuc_id)
    return await export_phieu_danh_gia_quy(quy, nam, db, target_cc)


@router.get("/bang-ke-cong-viec-quy/{quy}/{nam}/cua-cc/{cong_chuc_id}")
async def export_bang_ke_cong_viec_quy_cua_cc(
    quy: int,
    nam: int,
    cong_chuc_id: UUID,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """TDV/CCT tải bảng kê công việc quý (PL-02) của CC cụ thể."""
    target_cc = await _load_cc_cho_in_boi_tdv(db, current_user, cong_chuc_id)
    return await export_bang_ke_cong_viec_quy(quy, nam, db, target_cc)


@router.get("/phieu-danh-gia/{thang}/{nam}/cua-cc/{cong_chuc_id}")
async def export_phieu_danh_gia_thang_cua_cc(
    thang: int,
    nam: int,
    cong_chuc_id: UUID,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """TDV/CCT tải phiếu đánh giá tháng (PL-01A/01B) của CC cụ thể."""
    target_cc = await _load_cc_cho_in_boi_tdv(db, current_user, cong_chuc_id)
    return await export_phieu_danh_gia(thang, nam, db, target_cc)


@router.get("/bang-ke-cong-viec/{thang}/{nam}/cua-cc/{cong_chuc_id}")
async def export_bang_ke_cong_viec_thang_cua_cc(
    thang: int,
    nam: int,
    cong_chuc_id: UUID,
    db: DatabaseDep,
    current_user: ActiveUserDep,
):
    """TDV/CCT tải bảng kê công việc tháng (PL-02) của CC cụ thể."""
    target_cc = await _load_cc_cho_in_boi_tdv(db, current_user, cong_chuc_id)
    return await export_bang_ke_cong_viec(thang, nam, db, target_cc)
