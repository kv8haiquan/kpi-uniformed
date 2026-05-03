"""
bien_ban_service.py
=====================
Module 9 — Biên bản + Mock CKS + DOCX/PDF export.

Mock CKS (MVP):
- SHA-256 của noi_dung_json (sort_keys để deterministic)
- QR link verify: https://kv08.vn/verify/{hash}
- is_mock_signed=TRUE, qr_xac_thuc=URL, hash_noi_dung=SHA256
- KHÔNG implement PAdES (Phase 6)

Export:
- DOCX: python-docx render placeholder từ template default
- PDF: ReportLab Canvas (DejaVu Sans Unicode) + watermark "MOCK CKS" + QR
"""

from __future__ import annotations

import hashlib
import io
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from uuid import UUID, uuid4

from fastapi import HTTPException
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from meeting_service.config import settings
from meeting_service.models.bien_ban import BienBan
from meeting_service.models.cuoc_hop import CuocHop
from meeting_service.models.diem_danh import DiemDanh
from meeting_service.models.ket_luan import KetLuan
from meeting_service.models.mau_bieu import MauBieu
from meeting_service.models.thanh_phan import ThanhPhan
from meeting_service.models.y_kien import YKien
from meeting_service.services.audit_log_service import ghi_audit
from meeting_service.services.notification_service import gui_thong_bao
from shared.auth import TokenPayload


VERIFY_URL_BASE = "https://kv08.vn/verify"
FOLDER_BIEN_BAN = "bien-ban"


class BienBanService:
    def __init__(self, db: AsyncSession):
        self.db = db

    # ──────────────────────────────────────────────────────────────────
    # GET / CREATE-IF-EMPTY
    # ──────────────────────────────────────────────────────────────────
    async def get_or_init(
        self, cuoc_hop_id: UUID, user: TokenPayload
    ) -> BienBan:
        """Đọc biên bản. Nếu chưa có row → tạo mới với auto-fill."""
        result = await self.db.execute(
            select(BienBan).where(BienBan.cuoc_hop_id == cuoc_hop_id)
        )
        bb = result.scalar_one_or_none()
        if bb is not None:
            return bb

        # Auto-fill từ thành phần / điểm danh / ý kiến / kết luận
        auto_fill = await self._build_auto_fill(cuoc_hop_id)

        bb = BienBan(
            cuoc_hop_id=cuoc_hop_id,
            noi_dung_json=auto_fill,
            noi_dung_html=None,
            trang_thai="DANG_SOAN",
            nguoi_soan_id=UUID(user.sub),
        )
        self.db.add(bb)
        await self.db.flush()
        return bb

    async def _build_auto_fill(self, cuoc_hop_id: UUID) -> dict[str, Any]:
        """Snapshot dữ liệu cuộc họp lúc khởi tạo biên bản."""
        # Cuộc họp
        ch_res = await self.db.execute(
            select(CuocHop).where(CuocHop.id == cuoc_hop_id)
        )
        ch = ch_res.scalar_one_or_none()
        if ch is None:
            raise HTTPException(
                status_code=404,
                detail={"success": False, "error": {"code": "MEETING_NOT_FOUND",
                        "message": "Không tìm thấy cuộc họp"}},
            )

        # Thành phần
        tp_res = await self.db.execute(
            select(ThanhPhan).where(ThanhPhan.cuoc_hop_id == cuoc_hop_id)
        )
        tp_count = len(list(tp_res.scalars().all()))

        # Điểm danh summary
        dd_res = await self.db.execute(
            select(DiemDanh.trang_thai, func.count(DiemDanh.id))
            .where(DiemDanh.cuoc_hop_id == cuoc_hop_id)
            .group_by(DiemDanh.trang_thai)
        )
        dd_summary = {row[0]: row[1] for row in dd_res.fetchall()}

        # Ý kiến
        yk_res = await self.db.execute(
            select(YKien)
            .where(YKien.cuoc_hop_id == cuoc_hop_id, YKien.is_deleted.is_(False))
            .order_by(YKien.created_at.asc())
        )
        y_kien = [
            {"loai": yk.loai, "noi_dung": yk.noi_dung}
            for yk in yk_res.scalars().all()
        ]

        # Kết luận
        kl_res = await self.db.execute(
            select(KetLuan)
            .where(KetLuan.cuoc_hop_id == cuoc_hop_id, KetLuan.is_deleted.is_(False))
            .order_by(KetLuan.created_at.asc())
        )
        ket_luan = [
            {
                "noi_dung": kl.noi_dung,
                "han_hoan_thanh": kl.han_hoan_thanh.isoformat() if kl.han_hoan_thanh else None,
                "muc_uu_tien": kl.muc_uu_tien,
            }
            for kl in kl_res.scalars().all()
        ]

        return {
            "tieu_de": ch.tieu_de,
            "khoi": ch.khoi,
            "ngay_hop": ch.ngay_hop.isoformat(),
            "gio_bat_dau": ch.gio_bat_dau.isoformat(),
            "gio_ket_thuc": ch.gio_ket_thuc.isoformat() if ch.gio_ket_thuc else None,
            "dia_diem": ch.dia_diem,
            "tong_thanh_phan": tp_count,
            "diem_danh_summary": dd_summary,
            "y_kien": y_kien,
            "ket_luan": ket_luan,
            "noi_dung_thao_luan": "",  # Thư ký tự điền sau
        }

    # ──────────────────────────────────────────────────────────────────
    # UPDATE — Thư ký lưu nội dung biên bản
    # ──────────────────────────────────────────────────────────────────
    async def update(
        self,
        cuoc_hop_id: UUID,
        noi_dung_json: dict[str, Any],
        noi_dung_html: Optional[str],
        user: TokenPayload,
    ) -> BienBan:
        bb = await self.get_or_init(cuoc_hop_id, user)
        if bb.trang_thai in ("DA_KY", "CONG_BO"):
            raise HTTPException(
                status_code=409,
                detail={"success": False, "error": {"code": "BIEN_BAN_LOCKED",
                        "message": "Biên bản đã ký, không sửa được"}},
            )

        bb.noi_dung_json = noi_dung_json
        bb.noi_dung_html = noi_dung_html
        bb.updated_at = datetime.now(timezone.utc)

        await ghi_audit(
            self.db,
            hanh_dong="UPDATE_MINUTES",
            nguoi_thuc_hien_id=UUID(user.sub),
            doi_tuong_loai="bien_ban",
            doi_tuong_id=bb.id,
            chi_tiet={"cuoc_hop_id": str(cuoc_hop_id)},
        )
        await self.db.flush()
        return bb

    # ──────────────────────────────────────────────────────────────────
    # TRINH KY
    # ──────────────────────────────────────────────────────────────────
    async def trinh_ky(self, bien_ban_id: UUID, user: TokenPayload) -> BienBan:
        bb = await self._get(bien_ban_id)
        if bb.trang_thai != "DANG_SOAN":
            raise HTTPException(
                status_code=409,
                detail={"success": False, "error": {"code": "BIEN_BAN_TRANG_THAI",
                        "message": f"Biên bản đang ở trạng thái {bb.trang_thai}, không trình ký được"}},
            )

        # Chỉ thư ký mới trình ký
        if bb.nguoi_soan_id != UUID(user.sub) and not (
            user.is_admin or user.vai_tro in ("SUPER_ADMIN", "ADMIN")
        ):
            raise HTTPException(
                status_code=403,
                detail={"success": False, "error": {"code": "NO_PERMISSION",
                        "message": "Chỉ người soạn biên bản (thư ký) mới được trình ký"}},
            )

        bb.trang_thai = "TRINH_KY"
        bb.updated_at = datetime.now(timezone.utc)

        # Notify chu_toa
        ch_res = await self.db.execute(
            select(CuocHop).where(CuocHop.id == bb.cuoc_hop_id)
        )
        ch = ch_res.scalar_one()
        await gui_thong_bao(
            self.db,
            nguoi_nhan_id=ch.chu_toa_id,
            tieu_de=f"Biên bản chờ ký: {ch.tieu_de}",
            noi_dung="Thư ký đã trình biên bản. Vui lòng xem xét và ký.",
            sub_loai="BIEN_BAN_TRINH_KY",
            link_url=f"/hop-khong-giay/chi-tiet/{ch.id}/bien-ban",
            doi_tuong_id=bb.id,
            muc_do="QUAN_TRONG",
        )

        await ghi_audit(
            self.db,
            hanh_dong="SUBMIT_FOR_SIGN",
            nguoi_thuc_hien_id=UUID(user.sub),
            doi_tuong_loai="bien_ban",
            doi_tuong_id=bb.id,
            chi_tiet={"cuoc_hop_id": str(bb.cuoc_hop_id)},
        )
        await self.db.flush()
        return bb

    # ──────────────────────────────────────────────────────────────────
    # KY (Mock CKS)
    # ──────────────────────────────────────────────────────────────────
    async def ky_mock(self, bien_ban_id: UUID, user: TokenPayload) -> BienBan:
        bb = await self._get(bien_ban_id)
        if bb.trang_thai not in ("TRINH_KY", "DANG_SOAN"):
            raise HTTPException(
                status_code=409,
                detail={"success": False, "error": {"code": "BIEN_BAN_TRANG_THAI",
                        "message": f"Trạng thái {bb.trang_thai} không ký được"}},
            )

        # Chỉ chu_toa mới ký
        ch_res = await self.db.execute(
            select(CuocHop).where(CuocHop.id == bb.cuoc_hop_id)
        )
        ch = ch_res.scalar_one()
        user_id = UUID(user.sub)
        if ch.chu_toa_id != user_id and not (
            user.is_admin or user.vai_tro in ("SUPER_ADMIN", "ADMIN")
        ):
            raise HTTPException(
                status_code=403,
                detail={"success": False, "error": {"code": "NO_PERMISSION",
                        "message": "Chỉ chủ tọa được ký biên bản"}},
            )

        # Mock CKS: SHA-256 của noi_dung_json
        hash_value = compute_hash(bb.noi_dung_json or {})
        bb.hash_noi_dung = hash_value
        bb.qr_xac_thuc = f"{VERIFY_URL_BASE}/{hash_value}"
        bb.is_mock_signed = True
        bb.nguoi_ky_id = user_id
        bb.thoi_gian_ky = datetime.now(timezone.utc)
        bb.trang_thai = "DA_KY"
        bb.updated_at = datetime.now(timezone.utc)

        await ghi_audit(
            self.db,
            hanh_dong="SIGN_MINUTES",
            nguoi_thuc_hien_id=user_id,
            doi_tuong_loai="bien_ban",
            doi_tuong_id=bb.id,
            chi_tiet={
                "cuoc_hop_id": str(bb.cuoc_hop_id),
                "hash_noi_dung": hash_value,
                "is_mock_signed": True,
            },
        )

        # Notify thư ký + thành phần
        await gui_thong_bao(
            self.db,
            nguoi_nhan_id=bb.nguoi_soan_id,
            tieu_de=f"Biên bản đã được ký: {ch.tieu_de}",
            noi_dung=f"Chủ tọa đã ký. Hash xác thực: {hash_value[:16]}...",
            sub_loai="BIEN_BAN_DA_KY",
            link_url=f"/hop-khong-giay/chi-tiet/{ch.id}/bien-ban",
            doi_tuong_id=bb.id,
        )

        await self.db.flush()
        return bb

    # ──────────────────────────────────────────────────────────────────
    # EXPORT DOCX
    # ──────────────────────────────────────────────────────────────────
    async def xuat_docx(self, bien_ban_id: UUID, user: TokenPayload) -> dict:
        bb = await self._get(bien_ban_id)
        ch = (await self.db.execute(
            select(CuocHop).where(CuocHop.id == bb.cuoc_hop_id)
        )).scalar_one()

        from docx import Document  # lazy import

        doc = Document()
        # Nội dung biên bản
        title = doc.add_heading("BIÊN BẢN HỌP", 0)
        title.alignment = 1  # center

        data = bb.noi_dung_json or {}
        doc.add_paragraph(f"Tiêu đề: {data.get('tieu_de', ch.tieu_de)}")
        doc.add_paragraph(f"Khối: {data.get('khoi', ch.khoi)}")
        doc.add_paragraph(f"Ngày họp: {data.get('ngay_hop', ch.ngay_hop)}")
        doc.add_paragraph(f"Giờ bắt đầu: {data.get('gio_bat_dau', ch.gio_bat_dau)}")
        doc.add_paragraph(f"Địa điểm: {data.get('dia_diem', ch.dia_diem) or '—'}")
        doc.add_paragraph(f"Tổng thành phần: {data.get('tong_thanh_phan', 0)}")

        # Điểm danh summary
        dd = data.get("diem_danh_summary") or {}
        if dd:
            doc.add_heading("Điểm danh", 1)
            for k, v in dd.items():
                doc.add_paragraph(f"  {k}: {v}", style="List Bullet")

        # Nội dung thảo luận
        doc.add_heading("Nội dung thảo luận", 1)
        doc.add_paragraph(data.get("noi_dung_thao_luan") or "(chưa nhập)")

        # Ý kiến
        y_kien = data.get("y_kien") or []
        if y_kien:
            doc.add_heading("Ý kiến", 1)
            for yk in y_kien:
                doc.add_paragraph(f"[{yk.get('loai')}] {yk.get('noi_dung')}",
                                  style="List Bullet")

        # Kết luận
        ket_luan = data.get("ket_luan") or []
        if ket_luan:
            doc.add_heading("Kết luận / Nhiệm vụ", 1)
            for kl in ket_luan:
                han = kl.get("han_hoan_thanh") or "—"
                doc.add_paragraph(
                    f"{kl.get('noi_dung')} (Hạn: {han}, Ưu tiên: {kl.get('muc_uu_tien')})",
                    style="List Bullet",
                )

        # Mock CKS footer
        if bb.is_mock_signed:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("--- MOCK CKS (MVP) ---").bold = True
            doc.add_paragraph(f"Hash SHA-256: {bb.hash_noi_dung}")
            doc.add_paragraph(f"QR xác thực: {bb.qr_xac_thuc}")
            doc.add_paragraph(
                f"Thời gian ký: {bb.thoi_gian_ky.isoformat() if bb.thoi_gian_ky else '—'}"
            )

        # Save
        path = self._build_export_path(bb.cuoc_hop_id, "docx")
        doc.save(path)
        size = path.stat().st_size

        # Update DB
        rel_key = self._rel_key(path)
        bb.file_docx_minio_key = rel_key

        await ghi_audit(
            self.db,
            hanh_dong="EXPORT_MINUTES_DOCX",
            nguoi_thuc_hien_id=UUID(user.sub),
            doi_tuong_loai="bien_ban",
            doi_tuong_id=bb.id,
            chi_tiet={"file_size": size, "minio_key": rel_key},
        )
        await self.db.flush()

        return {
            "minio_key": rel_key,
            "file_size": size,
            "url_tai": f"/api/v1/hop-khong-giay/bien-ban/{bb.id}/file?dinh-dang=docx",
            "hash_noi_dung": bb.hash_noi_dung,
        }

    # ──────────────────────────────────────────────────────────────────
    # EXPORT PDF (ReportLab)
    # ──────────────────────────────────────────────────────────────────
    async def xuat_pdf(self, bien_ban_id: UUID, user: TokenPayload) -> dict:
        bb = await self._get(bien_ban_id)
        ch = (await self.db.execute(
            select(CuocHop).where(CuocHop.id == bb.cuoc_hop_id)
        )).scalar_one()

        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        import qrcode

        # Register Unicode fonts (idempotent — register lần thứ 2 không lỗi)
        font_dir = "/usr/share/fonts/truetype/dejavu"
        try:
            pdfmetrics.registerFont(TTFont("DejaVu", f"{font_dir}/DejaVuSans.ttf"))
            pdfmetrics.registerFont(TTFont("DejaVu-Bold", f"{font_dir}/DejaVuSans-Bold.ttf"))
        except Exception:
            pass  # đã register

        path = self._build_export_path(bb.cuoc_hop_id, "pdf")
        c = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4

        # ─ Watermark "MOCK CKS" diagonal 45° ─
        if bb.is_mock_signed:
            c.saveState()
            c.setFillGray(0.85)
            c.setFont("DejaVu-Bold", 60)
            c.translate(width / 2, height / 2)
            c.rotate(45)
            c.drawCentredString(0, 0, "MOCK CKS")
            c.restoreState()

        # ─ Title ─
        y = height - 2 * cm
        c.setFont("DejaVu-Bold", 18)
        c.drawCentredString(width / 2, y, "BIÊN BẢN HỌP")

        # ─ Body ─
        data = bb.noi_dung_json or {}
        c.setFont("DejaVu", 11)
        y -= 1.5 * cm
        for line in [
            f"Tiêu đề: {data.get('tieu_de', ch.tieu_de)}",
            f"Khối: {data.get('khoi', ch.khoi)}",
            f"Ngày họp: {data.get('ngay_hop', ch.ngay_hop)}",
            f"Giờ bắt đầu: {data.get('gio_bat_dau', ch.gio_bat_dau)}",
            f"Địa điểm: {data.get('dia_diem', ch.dia_diem) or '—'}",
            f"Tổng thành phần: {data.get('tong_thanh_phan', 0)}",
        ]:
            c.drawString(2 * cm, y, line)
            y -= 0.7 * cm

        # Điểm danh
        dd = data.get("diem_danh_summary") or {}
        if dd:
            y -= 0.3 * cm
            c.setFont("DejaVu-Bold", 12)
            c.drawString(2 * cm, y, "Điểm danh:")
            y -= 0.7 * cm
            c.setFont("DejaVu", 11)
            for k, v in dd.items():
                c.drawString(2.5 * cm, y, f"• {k}: {v}")
                y -= 0.6 * cm

        # Nội dung thảo luận
        y -= 0.3 * cm
        c.setFont("DejaVu-Bold", 12)
        c.drawString(2 * cm, y, "Nội dung thảo luận:")
        y -= 0.7 * cm
        c.setFont("DejaVu", 11)
        ndtl = data.get("noi_dung_thao_luan") or "(chưa nhập)"
        # Truncate dài → cắt ở 80 ký tự cho đơn giản MVP
        for chunk in [ndtl[i:i + 80] for i in range(0, len(ndtl), 80)][:5]:
            c.drawString(2.5 * cm, y, chunk)
            y -= 0.55 * cm

        # ─ QR góc dưới phải (nếu đã ký) ─
        if bb.is_mock_signed and bb.qr_xac_thuc:
            qr_img = qrcode.make(bb.qr_xac_thuc)
            qr_buf = io.BytesIO()
            qr_img.save(qr_buf, format="PNG")
            qr_buf.seek(0)
            from reportlab.lib.utils import ImageReader
            qr_size = 3 * cm
            c.drawImage(
                ImageReader(qr_buf),
                width - qr_size - 1.5 * cm,
                1.5 * cm,
                qr_size,
                qr_size,
            )

        # ─ Footer hash + thời gian ký ─
        if bb.is_mock_signed:
            c.setFont("DejaVu", 8)
            c.setFillGray(0.4)
            c.drawString(
                1.5 * cm, 1 * cm,
                f"Hash SHA-256: {bb.hash_noi_dung}",
            )
            c.drawString(
                1.5 * cm, 0.6 * cm,
                f"Thời gian ký (UTC): {bb.thoi_gian_ky.isoformat() if bb.thoi_gian_ky else '—'}",
            )

        c.showPage()
        c.save()

        size = path.stat().st_size
        rel_key = self._rel_key(path)
        bb.file_pdf_minio_key = rel_key

        await ghi_audit(
            self.db,
            hanh_dong="EXPORT_MINUTES_PDF",
            nguoi_thuc_hien_id=UUID(user.sub),
            doi_tuong_loai="bien_ban",
            doi_tuong_id=bb.id,
            chi_tiet={"file_size": size, "minio_key": rel_key},
        )
        await self.db.flush()

        return {
            "minio_key": rel_key,
            "file_size": size,
            "url_tai": f"/api/v1/hop-khong-giay/bien-ban/{bb.id}/file?dinh-dang=pdf",
            "hash_noi_dung": bb.hash_noi_dung,
        }

    # ──────────────────────────────────────────────────────────────────
    # HELPERS
    # ──────────────────────────────────────────────────────────────────
    async def _get(self, bien_ban_id: UUID) -> BienBan:
        result = await self.db.execute(
            select(BienBan).where(BienBan.id == bien_ban_id)
        )
        bb = result.scalar_one_or_none()
        if bb is None:
            raise HTTPException(
                status_code=404,
                detail={"success": False, "error": {"code": "BIEN_BAN_NOT_FOUND",
                        "message": "Không tìm thấy biên bản"}},
            )
        return bb

    @staticmethod
    def _build_export_path(cuoc_hop_id: UUID, ext: str) -> Path:
        root = Path(settings.upload_dir) / FOLDER_BIEN_BAN / str(cuoc_hop_id)
        root.mkdir(parents=True, exist_ok=True)
        return root / f"{uuid4().hex}_bien-ban.{ext}"

    @staticmethod
    def _rel_key(full_path: Path) -> str:
        """Convert full path → relative key trong bucket."""
        bucket = Path(settings.upload_dir).resolve()
        return str(full_path.resolve().relative_to(bucket))


# ════════════════════════════════════════════════════════════════════
# Pure functions — testable
# ════════════════════════════════════════════════════════════════════

def compute_hash(noi_dung_json: dict[str, Any]) -> str:
    """SHA-256 của JSON.dumps(sort_keys=True). Deterministic."""
    canonical = json.dumps(noi_dung_json, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()
