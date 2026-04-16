#!/usr/bin/env python3
"""
Script sửa lỗi paragraph trống trong template PL-Mẫu số 01A-CC.docx

Vấn đề: Sau paragraph "Vị trí việc làm..." có 1 paragraph chỉ chứa tab '\t'
        → gây ra dòng trống lửng lơ khi đã replace placeholder

Giải pháp: Xóa paragraph [7] (chỉ chứa '\t')
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement

# Add backend to path để import modules
sys.path.insert(0, str(Path(__file__).parent.parent))

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
TEMPLATE_FILE = "PL-Mẫu số 01A-CC.docx"


def main():
    template_path = TEMPLATES_DIR / TEMPLATE_FILE

    if not template_path.exists():
        print(f"❌ Template không tồn tại: {template_path}")
        return 1

    print(f"📄 Đọc template: {TEMPLATE_FILE}")
    doc = Document(template_path)

    # Backup
    backup_path = TEMPLATES_DIR / f"{TEMPLATE_FILE}.backup"
    doc.save(backup_path)
    print(f"💾 Đã backup: {backup_path}")

    # Dump paragraphs trước khi sửa
    print("\n📋 TRƯỚC KHI SỬA (paragraphs 5-10):")
    for i in range(5, min(10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        print(f"  [{i}] {repr(p.text)}")

    # Tìm paragraph cần xóa
    # Paragraph [6] là "Vị trí việc làm..."
    # Paragraph [7] là '\t' (cần xóa)
    # Paragraph [8] là "Đơn vị công tác..."

    target_index = None
    for i in range(5, min(10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        # Tìm paragraph chỉ chứa whitespace/tab ngay sau "Vị trí việc làm"
        if i > 0 and "Vị trí việc làm" in doc.paragraphs[i-1].text:
            if p.text.strip() == "":
                target_index = i
                print(f"\n🎯 Tìm thấy paragraph trống tại index {i}: {repr(p.text)}")
                break

    if target_index is None:
        print("⚠️  Không tìm thấy paragraph trống cần xóa")
        return 0

    # Xóa paragraph
    # Trong python-docx, xóa paragraph bằng cách xóa element khỏi parent
    p_element = doc.paragraphs[target_index]._element
    parent = p_element.getparent()
    parent.remove(p_element)

    print(f"✅ Đã xóa paragraph [{target_index}]")

    # Dump paragraphs sau khi sửa
    print("\n📋 SAU KHI SỬA (paragraphs 5-10):")
    # Re-load để lấy index mới
    for i in range(5, min(10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        print(f"  [{i}] {repr(p.text)}")

    # Save
    doc.save(template_path)
    print(f"\n💾 Đã lưu template: {template_path}")
    print("\n✅ HOÀN THÀNH")
    return 0


if __name__ == "__main__":
    sys.exit(main())
