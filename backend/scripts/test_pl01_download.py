#!/usr/bin/env python3
"""
Script test download và verify PL-01A/PL-01B output.
Verify header không còn dòng trống giữa "Vị trí việc làm" và "Đơn vị công tác".
"""

import sys
import asyncio
from pathlib import Path
from docx import Document

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.db.session import AsyncSessionLocal
from app.api.v1.endpoints.in_bang_ke import tao_phieu_danh_gia_pl01
from io import BytesIO


async def test_download(ma_cc: str, thang: int, nam: int, output_file: str):
    """
    Test download và dump paragraphs header.
    """
    print(f"\n{'='*60}")
    print(f"TEST: ma_cc={ma_cc}, thang={thang}/{nam}")
    print(f"{'='*60}")

    async with AsyncSessionLocal() as db:
        try:
            # Gọi hàm tạo phiếu
            file_bytes = await tao_phieu_danh_gia_pl01(db, ma_cc, thang, nam)

            # Lưu file
            with open(output_file, "wb") as f:
                f.write(file_bytes)

            print(f"✅ Đã lưu file: {output_file} ({len(file_bytes)} bytes)")

            # Load lại và dump paragraphs header
            doc = Document(BytesIO(file_bytes))
            print("\n📋 HEADER (paragraphs 5-10):")
            for i in range(5, min(10, len(doc.paragraphs))):
                p = doc.paragraphs[i]
                print(f"  [{i}] {repr(p.text)}")

            # Check có dòng trống giữa "Vị trí việc làm" và "Đơn vị công tác" không
            found_vi_tri = False
            found_don_vi = False
            has_empty_between = False

            for i, p in enumerate(doc.paragraphs[:15]):
                text = p.text
                if "Vị trí việc làm" in text or "Chức vụ, chức danh" in text:
                    found_vi_tri = True
                    continue
                if found_vi_tri and not found_don_vi:
                    if "Đơn vị công tác" in text:
                        found_don_vi = True
                    elif text.strip() == "" or text == "\t":
                        has_empty_between = True
                        print(f"\n⚠️  PHÁT HIỆN DÒNG TRỐNG tại paragraph [{i}]: {repr(text)}")

            if has_empty_between:
                print("\n❌ LỖI: Vẫn còn dòng trống giữa 'Vị trí việc làm' và 'Đơn vị công tác'")
                return 1
            else:
                print("\n✅ OK: Không có dòng trống giữa 'Vị trí việc làm' và 'Đơn vị công tác'")
                return 0

        except Exception as e:
            print(f"❌ Lỗi: {e}")
            import traceback
            traceback.print_exc()
            return 1


async def main():
    # Test case 1: PL-01A (công chức thường — template đã sửa)
    result1 = await test_download(
        ma_cc="20ZZ-0447",
        thang=1,
        nam=2026,
        output_file="/tmp/test_pl01a.docx"
    )

    # Test case 2: PL-01B (lãnh đạo — template không bị lỗi)
    result2 = await test_download(
        ma_cc="20ZZ-0224",
        thang=2,
        nam=2026,
        output_file="/tmp/test_pl01b.docx"
    )

    if result1 == 0 and result2 == 0:
        print("\n" + "="*60)
        print("✅ TẤT CẢ TESTS PASS")
        print("="*60)
        return 0
    else:
        print("\n" + "="*60)
        print("❌ CÓ TEST FAILED")
        print("="*60)
        return 1


if __name__ == "__main__":
    exit_code = asyncio.run(main())
    sys.exit(exit_code)
