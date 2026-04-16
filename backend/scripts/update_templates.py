"""
Script cập nhật template DOCX với merge field cho tính năng in bảng kê.

PL-01A/01B: Phiếu đánh giá (thêm placeholder cho thông tin cá nhân, kỳ đánh giá, điểm)
PL-02: Bảng kê công việc (thêm placeholder cho thông tin + bảng công việc)

Run: python scripts/update_templates.py
"""

from docx import Document
from pathlib import Path
import re

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


def replace_text_in_paragraph(para, old_text, new_text):
    """
    Replace text trong paragraph, xử lý trường hợp text bị split qua nhiều runs.
    """
    if old_text in para.text:
        # Tìm vị trí text cần replace
        full_text = para.text
        if old_text not in full_text:
            return False

        # Merge all runs into one để dễ replace
        for run in para.runs:
            run.text = ""
        para.runs[0].text = full_text.replace(old_text, new_text)
        return True
    return False


def update_pl01a_pl01b(template_path: Path, is_lanh_dao: bool):
    """
    Cập nhật template PL-01A (CC) hoặc PL-01B (Lãnh đạo).

    Thêm merge field:
    - {{ky_danh_gia}} - Tháng/Năm
    - {{ho_ten}} - Họ tên CC
    - {{chuc_vu}} - Chức vụ (PL-01B) hoặc Vị trí việc làm (PL-01A)
    - {{don_vi}} - Đơn vị công tác
    - {{diem_tieu_chi}} - Điểm tiêu chí chung
    - {{diem_tong}} - Tổng điểm KPI
    """
    doc = Document(template_path)

    print(f"\n{'='*80}")
    print(f"Cập nhật {template_path.name}")
    print('='*80)

    # Duyệt paragraphs
    for para in doc.paragraphs:
        # Kỳ đánh giá
        if "Kỳ theo dõi, đánh giá:" in para.text:
            replace_text_in_paragraph(para, "Kỳ theo dõi, đánh giá:......", "Kỳ theo dõi, đánh giá: {{ky_danh_gia}}")
            print("  ✓ Thêm {{ky_danh_gia}}")

        # Họ tên
        if para.text.strip().startswith("Họ và tên:"):
            replace_text_in_paragraph(para, "Họ và tên:\t", "Họ và tên: {{ho_ten}}")
            print("  ✓ Thêm {{ho_ten}}")

        # Chức vụ (01B) hoặc Vị trí việc làm (01A)
        if is_lanh_dao and "Chức vụ, chức danh:" in para.text:
            replace_text_in_paragraph(para, "Chức vụ, chức danh:\t", "Chức vụ, chức danh: {{chuc_vu}}")
            print("  ✓ Thêm {{chuc_vu}}")
        elif not is_lanh_dao and "Vị trí việc làm/nhiệm vụ được phân công chuyên trách:" in para.text:
            replace_text_in_paragraph(para, "Vị trí việc làm/nhiệm vụ được phân công chuyên trách:\t",
                                     "Vị trí việc làm/nhiệm vụ được phân công chuyên trách: {{chuc_vu}}")
            print("  ✓ Thêm {{chuc_vu}} (vị trí việc làm)")

        # Đơn vị công tác
        if "Đơn vị công tác:" in para.text:
            replace_text_in_paragraph(para, "Đơn vị công tác:\t", "Đơn vị công tác: {{don_vi}}")
            print("  ✓ Thêm {{don_vi}}")

        # Điểm tiêu chí chung
        if "1. Điểm tiêu chí chung (tối đa 30 điểm):" in para.text:
            replace_text_in_paragraph(para, "1. Điểm tiêu chí chung (tối đa 30 điểm): \t",
                                     "1. Điểm tiêu chí chung (tối đa 30 điểm): {{diem_tieu_chi}}")
            print("  ✓ Thêm {{diem_tieu_chi}}")

    # Bảng tiêu chí: thêm điểm vào cột (4) - sẽ xử lý khi export
    # (Logic phức tạp hơn, sẽ điền trực tiếp trong endpoint export)

    # Lưu file
    output_path = template_path.parent / f"{template_path.stem}_v2{template_path.suffix}"
    doc.save(output_path)
    print(f"  → Saved: {output_path.name}")

    # Backup original và rename
    backup_path = template_path.parent / f"{template_path.stem}_original{template_path.suffix}"
    if not backup_path.exists():
        import shutil
        shutil.copy(template_path, backup_path)
        print(f"  → Backup: {backup_path.name}")

    # Replace original
    import shutil
    shutil.move(output_path, template_path)
    print(f"  → Replaced original: {template_path.name}")


def update_pl02(template_path: Path):
    """
    Cập nhật template PL-02 (Bảng kê công việc).

    Merge field:
    - {{ky_danh_gia}} - Tháng/Năm
    - {{tong_cv}} - Tổng số công việc
    - {{chua_ht}} - Chưa hoàn thành
    - {{da_ht}} - Đã hoàn thành
    - {{dung_han}} - Đúng tiến độ
    - {{ty_le_dung_han}} - % đúng hạn
    - {{tre_han}} - Chậm tiến độ
    - {{ty_le_tre_han}} - % trễ hạn

    Bảng công việc: sẽ điền động trong endpoint export
    """
    doc = Document(template_path)

    print(f"\n{'='*80}")
    print(f"Cập nhật {template_path.name}")
    print('='*80)

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text

        if "Kỳ theo dõi, đánh giá:" in text:
            replace_text_in_paragraph(para, "Kỳ theo dõi, đánh giá:......", "Kỳ theo dõi, đánh giá: {{ky_danh_gia}}")
            print("  ✓ Thêm {{ky_danh_gia}}")

        if "Tổng số công việc cần thực hiện trong kỳ" in text:
            replace_text_in_paragraph(para, "Tổng số công việc cần thực hiện trong kỳ (cộng cột số (4)): \t",
                                     "Tổng số công việc cần thực hiện trong kỳ (cộng cột số (4)): {{tong_cv}}")
            print("  ✓ Thêm {{tong_cv}}")

        if "2. Tổng số công việc chưa hoàn thành trong kỳ" in text:
            replace_text_in_paragraph(para, "2. Tổng số công việc chưa hoàn thành trong kỳ (cộng cột (7)):: \t",
                                     "2. Tổng số công việc chưa hoàn thành trong kỳ (cộng cột (7)): {{chua_ht}}")
            print("  ✓ Thêm {{chua_ht}}")

        if "3. Tổng số công việc đã hoàn thành trong kỳ:" in text:
            replace_text_in_paragraph(para, "3. Tổng số công việc đã hoàn thành trong kỳ: \t",
                                     "3. Tổng số công việc đã hoàn thành trong kỳ: {{da_ht}}")
            print("  ✓ Thêm {{da_ht}}")

        if "4. Tổng số công việc hoàn thành đạt tiến độ trong kỳ" in text:
            replace_text_in_paragraph(para, "4. Tổng số công việc hoàn thành đạt tiến độ trong kỳ (cộng cột số (8))\t",
                                     "4. Tổng số công việc hoàn thành đạt tiến độ trong kỳ (cộng cột số (8)): {{dung_han}}")
            print("  ✓ Thêm {{dung_han}}")

        if "Tỷ lệ: ………..%" in text and "tiến độ/Tổng số công việc đã hoàn thành" in text:
            if "chậm tiến độ" not in text:  # Tỷ lệ đúng hạn
                replace_text_in_paragraph(para, "Tỷ lệ: ………..%", "Tỷ lệ: {{ty_le_dung_han}}%")
                print("  ✓ Thêm {{ty_le_dung_han}}")

        if "5. Tổng số công việc hoàn thành chậm tiến độ trong tháng" in text:
            replace_text_in_paragraph(para, "5. Tổng số công việc hoàn thành chậm tiến độ trong tháng (cộng cột số (9))\t",
                                     "5. Tổng số công việc hoàn thành chậm tiến độ trong tháng (cộng cột số (9)): {{tre_han}}")
            print("  ✓ Thêm {{tre_han}}")

        if "Tỷ lệ: ………..%" in text and "chậm tiến độ/Tổng số công việc đã hoàn thành" in text:
            replace_text_in_paragraph(para, "Tỷ lệ: ………..%", "Tỷ lệ: {{ty_le_tre_han}}%")
            print("  ✓ Thêm {{ty_le_tre_han}}")

    # Lưu
    output_path = template_path.parent / f"{template_path.stem}_v2{template_path.suffix}"
    doc.save(output_path)
    print(f"  → Saved: {output_path.name}")

    # Backup và replace
    backup_path = template_path.parent / f"{template_path.stem}_original{template_path.suffix}"
    if not backup_path.exists():
        import shutil
        shutil.copy(template_path, backup_path)
        print(f"  → Backup: {backup_path.name}")

    import shutil
    shutil.move(output_path, template_path)
    print(f"  → Replaced original: {template_path.name}")


if __name__ == "__main__":
    print("Updating templates with merge fields...")

    # PL-01A (CC không lãnh đạo)
    pl01a = TEMPLATES_DIR / "PL-Mẫu số 01A-CC.docx"
    if pl01a.exists():
        update_pl01a_pl01b(pl01a, is_lanh_dao=False)

    # PL-01B (Lãnh đạo)
    pl01b = TEMPLATES_DIR / "PL-Mẫu số 01B-LĐ.docx"
    if pl01b.exists():
        update_pl01a_pl01b(pl01b, is_lanh_dao=True)

    # PL-02 (Bảng kê)
    pl02 = TEMPLATES_DIR / "PL-Mẫu số 02-Bang ke.docx"
    if pl02.exists():
        update_pl02(pl02)

    print("\n✅ DONE. Templates updated.")
